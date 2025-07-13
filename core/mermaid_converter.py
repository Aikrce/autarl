#!/usr/bin/env python3
"""
Mermaid Diagram to Word Converter
支持将Markdown中的Mermaid图表转换为Word文档中的图片
"""

import re
import os
import tempfile
import subprocess
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import base64
import requests
from PIL import Image
import io


class MermaidConverter:
    """Mermaid图表转换器"""
    
    def __init__(self, method='api'):
        """
        初始化Mermaid转换器
        
        Args:
            method: 转换方法 ('api', 'cli', 'web')
                - api: 使用在线API服务
                - cli: 使用本地mermaid-cli
                - web: 使用Web渲染服务
        """
        self.method = method
        self.temp_dir = tempfile.mkdtemp()
        
    def __del__(self):
        """清理临时文件"""
        import shutil
        if hasattr(self, 'temp_dir') and os.path.exists(self.temp_dir):
            shutil.rmtree(self.temp_dir)
    
    def extract_mermaid_blocks(self, markdown_content):
        """
        从Markdown内容中提取Mermaid代码块
        
        Returns:
            列表，每个元素包含 (start_pos, end_pos, mermaid_code)
        """
        mermaid_blocks = []
        
        # 匹配 ```mermaid 代码块
        pattern = r'```mermaid\s*\n(.*?)```'
        matches = re.finditer(pattern, markdown_content, re.DOTALL)
        
        for match in matches:
            start_pos = match.start()
            end_pos = match.end()
            mermaid_code = match.group(1).strip()
            mermaid_blocks.append({
                'start': start_pos,
                'end': end_pos,
                'code': mermaid_code,
                'type': self._detect_diagram_type(mermaid_code)
            })
        
        return mermaid_blocks
    
    def _detect_diagram_type(self, mermaid_code):
        """检测Mermaid图表类型"""
        first_line = mermaid_code.split('\n')[0].strip().lower()
        
        diagram_types = {
            'graph': 'flowchart',
            'flowchart': 'flowchart',
            'sequencediagram': 'sequence',
            'sequence': 'sequence',
            'classDiagram': 'class',
            'class': 'class',
            'stateDiagram': 'state',
            'state': 'state',
            'erDiagram': 'er',
            'er': 'er',
            'gantt': 'gantt',
            'pie': 'pie',
            'journey': 'journey',
            'gitgraph': 'git'
        }
        
        for key, value in diagram_types.items():
            if first_line.startswith(key.lower()):
                return value
        
        return 'flowchart'  # 默认类型
    
    def convert_mermaid_to_image(self, mermaid_code, output_path=None):
        """
        将Mermaid代码转换为图片
        
        Args:
            mermaid_code: Mermaid图表代码
            output_path: 输出图片路径（可选）
            
        Returns:
            图片路径或图片数据
        """
        if self.method == 'api':
            return self._convert_via_api(mermaid_code, output_path)
        elif self.method == 'cli':
            return self._convert_via_cli(mermaid_code, output_path)
        elif self.method == 'web':
            return self._convert_via_web(mermaid_code, output_path)
        else:
            raise ValueError(f"不支持的转换方法: {self.method}")
    
    def _convert_via_api(self, mermaid_code, output_path=None):
        """通过在线API转换（使用mermaid.ink服务）"""
        # 对Mermaid代码进行Base64编码
        encoded = base64.urlsafe_b64encode(mermaid_code.encode()).decode()
        
        # 使用mermaid.ink服务
        url = f"https://mermaid.ink/img/{encoded}"
        
        try:
            response = requests.get(url, timeout=30)
            response.raise_for_status()
            
            if output_path:
                with open(output_path, 'wb') as f:
                    f.write(response.content)
                return output_path
            else:
                return response.content
                
        except Exception as e:
            print(f"API转换失败: {e}")
            # 回退到本地渲染
            return self._convert_via_svg_placeholder(mermaid_code, output_path)
    
    def _convert_via_cli(self, mermaid_code, output_path=None):
        """通过本地mermaid-cli转换"""
        if not output_path:
            output_path = os.path.join(self.temp_dir, f"mermaid_{hash(mermaid_code)}.png")
        
        # 创建临时mermaid文件
        mermaid_file = os.path.join(self.temp_dir, "temp.mmd")
        with open(mermaid_file, 'w', encoding='utf-8') as f:
            f.write(mermaid_code)
        
        try:
            # 使用mmdc命令（需要先安装 npm install -g @mermaid-js/mermaid-cli）
            subprocess.run([
                'mmdc',
                '-i', mermaid_file,
                '-o', output_path,
                '-t', 'default',
                '-w', '800',
                '-H', '600'
            ], check=True, capture_output=True)
            
            return output_path
            
        except (subprocess.CalledProcessError, FileNotFoundError) as e:
            print(f"mermaid-cli未安装或执行失败: {e}")
            # 回退到占位符
            return self._convert_via_svg_placeholder(mermaid_code, output_path)
    
    def _convert_via_web(self, mermaid_code, output_path=None):
        """通过Web渲染服务转换（使用kroki.io）"""
        try:
            # 使用kroki.io服务
            url = "https://kroki.io/mermaid/png"
            
            response = requests.post(
                url,
                headers={'Content-Type': 'text/plain'},
                data=mermaid_code.encode('utf-8'),
                timeout=30
            )
            response.raise_for_status()
            
            if output_path:
                with open(output_path, 'wb') as f:
                    f.write(response.content)
                return output_path
            else:
                return response.content
                
        except Exception as e:
            print(f"Web服务转换失败: {e}")
            return self._convert_via_svg_placeholder(mermaid_code, output_path)
    
    def _convert_via_svg_placeholder(self, mermaid_code, output_path=None):
        """创建一个占位符图片（当其他方法失败时）"""
        from PIL import Image, ImageDraw, ImageFont
        
        # 创建占位符图片
        width, height = 800, 400
        img = Image.new('RGB', (width, height), color='white')
        draw = ImageDraw.Draw(img)
        
        # 绘制边框
        draw.rectangle([(0, 0), (width-1, height-1)], outline='black', width=2)
        
        # 添加文本
        title = "Mermaid Diagram"
        subtitle = f"Type: {self._detect_diagram_type(mermaid_code)}"
        
        try:
            # 尝试使用系统字体
            title_font = ImageFont.truetype("Arial.ttf", 24)
            text_font = ImageFont.truetype("Arial.ttf", 14)
        except:
            # 使用默认字体
            title_font = ImageFont.load_default()
            text_font = ImageFont.load_default()
        
        # 绘制标题
        title_bbox = draw.textbbox((0, 0), title, font=title_font)
        title_width = title_bbox[2] - title_bbox[0]
        title_height = title_bbox[3] - title_bbox[1]
        draw.text(((width - title_width) // 2, 20), title, fill='black', font=title_font)
        
        # 绘制副标题
        subtitle_bbox = draw.textbbox((0, 0), subtitle, font=text_font)
        subtitle_width = subtitle_bbox[2] - subtitle_bbox[0]
        draw.text(((width - subtitle_width) // 2, 60), subtitle, fill='gray', font=text_font)
        
        # 添加代码预览（前几行）
        code_lines = mermaid_code.split('\n')[:5]
        y_offset = 100
        for line in code_lines:
            if len(line) > 50:
                line = line[:50] + '...'
            draw.text((20, y_offset), line, fill='black', font=text_font)
            y_offset += 20
        
        # 保存图片
        if output_path:
            img.save(output_path, 'PNG')
            return output_path
        else:
            buffer = io.BytesIO()
            img.save(buffer, 'PNG')
            return buffer.getvalue()
    
    def add_mermaid_to_document(self, doc, mermaid_code, width=6.0):
        """
        将Mermaid图表添加到Word文档
        
        Args:
            doc: python-docx Document对象
            mermaid_code: Mermaid代码
            width: 图片宽度（英寸）
        """
        # 转换为图片
        image_path = os.path.join(self.temp_dir, f"mermaid_{hash(mermaid_code)}.png")
        self.convert_mermaid_to_image(mermaid_code, image_path)
        
        # 添加图片到文档
        if os.path.exists(image_path):
            paragraph = doc.add_paragraph()
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            run = paragraph.add_run()
            run.add_picture(image_path, width=Inches(width))
            
            # 添加图表说明
            caption = doc.add_paragraph()
            caption.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            caption_run = caption.add_run(f"图: {self._detect_diagram_type(mermaid_code).title()}图表")
            caption_run.font.size = Pt(10)
            caption_run.font.italic = True
            
            # 添加空行
            doc.add_paragraph()


class EnhancedMarkdownConverter:
    """增强的Markdown转换器，支持Mermaid图表"""
    
    def __init__(self, mermaid_method='api'):
        self.mermaid_converter = MermaidConverter(method=mermaid_method)
        
    def convert_with_mermaid(self, markdown_file, output_file):
        """转换包含Mermaid图表的Markdown文件"""
        # 读取Markdown文件
        with open(markdown_file, 'r', encoding='utf-8') as f:
            content = f.read()
        
        # 创建Word文档
        doc = Document()
        
        # 提取Mermaid块
        mermaid_blocks = self.mermaid_converter.extract_mermaid_blocks(content)
        
        # 处理内容
        current_pos = 0
        
        for block in mermaid_blocks:
            # 处理Mermaid块之前的内容
            before_content = content[current_pos:block['start']]
            self._process_markdown_content(doc, before_content)
            
            # 处理Mermaid图表
            self.mermaid_converter.add_mermaid_to_document(doc, block['code'])
            
            current_pos = block['end']
        
        # 处理剩余内容
        remaining_content = content[current_pos:]
        self._process_markdown_content(doc, remaining_content)
        
        # 保存文档
        doc.save(output_file)
        print(f"✓ 成功转换文档（含Mermaid图表）: {output_file}")
        
    def _process_markdown_content(self, doc, content):
        """处理普通Markdown内容"""
        lines = content.split('\n')
        
        for line in lines:
            line = line.strip()
            if not line:
                doc.add_paragraph()
                continue
                
            if line.startswith('# '):
                doc.add_heading(line[2:], level=1)
            elif line.startswith('## '):
                doc.add_heading(line[3:], level=2)
            elif line.startswith('### '):
                doc.add_heading(line[4:], level=3)
            elif line.startswith('- ') or line.startswith('* '):
                doc.add_paragraph(line[2:], style='List Bullet')
            elif re.match(r'^\d+\.\s', line):
                doc.add_paragraph(line[3:], style='List Number')
            else:
                doc.add_paragraph(line)


def main():
    """主函数 - 演示用法"""
    # 创建转换器（使用Web渲染服务，避免SSL问题）
    converter = EnhancedMarkdownConverter(mermaid_method='web')
    
    # 测试文件
    test_files = [
        ('mermaid_test.md', 'mermaid_test_output.docx')
    ]
    
    for input_file, output_file in test_files:
        if os.path.exists(input_file):
            try:
                print(f"正在转换 {input_file}...")
                converter.convert_with_mermaid(input_file, output_file)
            except Exception as e:
                print(f"转换失败: {e}")
        else:
            print(f"测试文件 {input_file} 不存在")


if __name__ == "__main__":
    main()