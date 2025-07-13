#!/usr/bin/env python3
"""
修复版的完整Markdown转换器，正确处理Mermaid图表
"""

import re
import os
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# 导入现有的转换器
from enhanced_table_converter import AdvancedTableConverter
from mermaid_converter import MermaidConverter


class FixedCompleteMarkdownConverter:
    """修复版的完整Markdown转换器"""
    
    def __init__(self, mermaid_method='web'):
        self.table_converter = AdvancedTableConverter()
        self.mermaid_converter = MermaidConverter(method=mermaid_method)
        
    def convert(self, markdown_file, output_file):
        """转换Markdown文件到Word"""
        # 读取文件
        with open(markdown_file, 'r', encoding='utf-8') as f:
            content = f.read()
        
        # 创建Word文档
        doc = Document()
        
        # 处理内容
        self._process_content(content, doc)
        
        # 保存文档
        doc.save(output_file)
        print(f"✓ 成功转换文档（支持表格和Mermaid图表）: {output_file}")
        return True
    
    def _process_content(self, content, doc):
        """处理Markdown内容"""
        lines = content.split('\n')
        i = 0
        
        while i < len(lines):
            line = lines[i]
            
            # 检查是否是Mermaid代码块的开始
            if line.strip() == '```mermaid':
                # 收集Mermaid代码
                mermaid_lines = []
                i += 1
                while i < len(lines) and lines[i].strip() != '```':
                    mermaid_lines.append(lines[i])
                    i += 1
                
                # 转换Mermaid图表
                if mermaid_lines:
                    mermaid_code = '\n'.join(mermaid_lines)
                    print(f"发现Mermaid图表，类型: {self.mermaid_converter._detect_diagram_type(mermaid_code)}")
                    self.mermaid_converter.add_mermaid_to_document(doc, mermaid_code)
                
                # 跳过结束标记
                i += 1
                continue
            
            # 检查是否是表格
            if '|' in line and not line.strip().startswith('```'):
                # 收集表格数据
                table_lines = []
                table_start = i
                
                while i < len(lines) and '|' in lines[i].strip() and not lines[i].strip().startswith('```'):
                    table_lines.append(lines[i])
                    i += 1
                
                # 处理表格
                if table_lines:
                    # 解析表格数据
                    table_data = []
                    for table_line in table_lines:
                        if not re.match(r'^\|[\s\-\|:]+\|?$', table_line.strip()):
                            cells = [cell.strip() for cell in table_line.split('|')]
                            if cells and cells[0] == '':
                                cells = cells[1:]
                            if cells and cells[-1] == '':
                                cells = cells[:-1]
                            if cells:
                                table_data.append(cells)
                    
                    if table_data:
                        self.table_converter.add_table_to_document(doc, {'data': table_data})
                
                continue
            
            # 处理其他Markdown元素
            self._process_normal_line(doc, line)
            i += 1
    
    def _process_normal_line(self, doc, line):
        """处理普通Markdown行"""
        # 分隔线
        if line.strip() in ['---', '***', '___']:
            doc.add_page_break()
            return
        
        # 空行
        if not line.strip():
            doc.add_paragraph()
            return
        
        # 标题
        if line.startswith('#'):
            level = len(line) - len(line.lstrip('#'))
            title_text = line.lstrip('#').strip()
            if title_text:  # 只有当标题有实际内容时才添加
                doc.add_heading(title_text, level=min(level, 6))
            return
        
        # 列表
        if line.strip().startswith(('- ', '* ', '+ ')):
            doc.add_paragraph(line.strip()[2:], style='List Bullet')
            return
        
        if re.match(r'^\d+\.\s', line.strip()):
            text = re.sub(r'^\d+\.\s', '', line.strip())
            doc.add_paragraph(text, style='List Number')
            return
        
        # 引用
        if line.startswith('> '):
            quote_para = doc.add_paragraph(line[2:])
            quote_para.paragraph_format.left_indent = Inches(0.5)
            return
        
        # 普通段落
        para = doc.add_paragraph()
        self._process_inline_formatting(para, line)
    
    def _process_inline_formatting(self, paragraph, text):
        """处理内联格式"""
        # 简化的内联格式处理
        parts = re.split(r'(\*\*[^*]+\*\*|__[^_]+__|`[^`]+`)', text)
        
        for part in parts:
            if not part:
                continue
            
            if part.startswith('**') and part.endswith('**'):
                run = paragraph.add_run(part[2:-2])
                run.bold = True
            elif part.startswith('__') and part.endswith('__'):
                run = paragraph.add_run(part[2:-2])
                run.bold = True
            elif part.startswith('`') and part.endswith('`'):
                run = paragraph.add_run(part[1:-1])
                run.font.name = 'Consolas'
                run.font.size = Pt(10)
            else:
                paragraph.add_run(part)


def main():
    """测试修复版转换器"""
    import sys
    
    if len(sys.argv) < 2:
        print("使用方法: python fixed_complete_converter.py <input.md> [output.docx]")
        # 默认测试
        converter = FixedCompleteMarkdownConverter()
        converter.convert('test_mermaid_issue.md', 'test_mermaid_fixed.docx')
    else:
        input_file = sys.argv[1]
        output_file = sys.argv[2] if len(sys.argv) > 2 else input_file.replace('.md', '_converted.docx')
        
        converter = FixedCompleteMarkdownConverter()
        converter.convert(input_file, output_file)


if __name__ == "__main__":
    main()