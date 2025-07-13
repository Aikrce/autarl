#!/usr/bin/env python3
"""
Enhanced Markdown to Word converter with advanced table support
增强的Markdown到Word转换器，支持高级表格转换
"""

import re
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import markdown2
from bs4 import BeautifulSoup


class AdvancedTableConverter:
    """高级表格转换器"""
    
    def __init__(self):
        self.table_style = 'Table Grid'  # 默认表格样式
        
    def parse_markdown_table(self, markdown_text):
        """解析Markdown表格"""
        tables = []
        lines = markdown_text.split('\n')
        
        i = 0
        while i < len(lines):
            line = lines[i].strip()
            # 检测表格开始（包含 | 分隔符的行）
            if '|' in line and not line.startswith('```'):
                table_data = []
                table_start = i
                
                # 读取表格的所有行
                while i < len(lines) and '|' in lines[i].strip():
                    table_line = lines[i].strip()
                    if table_line:
                        # 跳过分隔行（如 |-----|-----|）
                        if not re.match(r'^\|[\s\-\|:]+\|?$', table_line):
                            # 解析表格行
                            cells = [cell.strip() for cell in table_line.split('|')]
                            # 移除首尾的空元素（由于开头结尾的|导致）
                            if cells and cells[0] == '':
                                cells = cells[1:]
                            if cells and cells[-1] == '':
                                cells = cells[:-1]
                            
                            if cells:  # 只添加非空行
                                table_data.append(cells)
                    i += 1
                
                if table_data:
                    tables.append({
                        'data': table_data,
                        'start_line': table_start,
                        'end_line': i - 1
                    })
                i -= 1  # 回退一行，因为外层循环会自增
            i += 1
            
        return tables
    
    def add_table_to_document(self, doc, table_data):
        """将表格数据添加到Word文档"""
        if not table_data or not table_data['data']:
            return
            
        rows = table_data['data']
        if not rows:
            return
            
        # 确定表格的行数和列数
        max_cols = max(len(row) for row in rows) if rows else 0
        if max_cols == 0:
            return
            
        # 创建表格
        table = doc.add_table(rows=len(rows), cols=max_cols)
        table.style = self.table_style
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # 填充表格数据
        for row_idx, row_data in enumerate(rows):
            for col_idx in range(max_cols):
                cell = table.cell(row_idx, col_idx)
                
                # 获取单元格内容
                if col_idx < len(row_data):
                    cell_content = row_data[col_idx].strip()
                else:
                    cell_content = ""  # 空单元格
                
                # 设置单元格内容
                if cell_content:
                    # 处理格式化文本（粗体、斜体等）
                    self._format_cell_content(cell, cell_content)
                
                # 设置表头样式（第一行）
                if row_idx == 0:
                    self._format_header_cell(cell)
                else:
                    self._format_data_cell(cell)
        
        # 设置表格边框
        self._set_table_borders(table)
        
        # 添加表格后的空行
        doc.add_paragraph()
    
    def _format_cell_content(self, cell, content):
        """格式化单元格内容"""
        paragraph = cell.paragraphs[0]
        paragraph.clear()
        
        # 处理Markdown格式
        # 粗体 **text** 或 __text__
        bold_pattern = r'\*\*(.*?)\*\*|__(.*?)__'
        # 斜体 *text* 或 _text_
        italic_pattern = r'\*(.*?)\*|(?<!_)_(.*?)_(?!_)'
        # 行内代码 `code`
        code_pattern = r'`(.*?)`'
        
        # 分割文本并应用格式
        parts = re.split(f'({bold_pattern}|{italic_pattern}|{code_pattern})', content)
        
        for part in parts:
            if not part:
                continue
                
            run = paragraph.add_run()
            
            # 检查格式类型
            if re.match(bold_pattern, part):
                # 粗体文本
                text = re.sub(r'\*\*(.*?)\*\*|__(.*?)__', r'\1\2', part)
                run.text = text
                run.bold = True
            elif re.match(italic_pattern, part):
                # 斜体文本
                text = re.sub(r'\*(.*?)\*|(?<!_)_(.*?)_(?!_)', r'\1\2', part)
                run.text = text
                run.italic = True
            elif re.match(code_pattern, part):
                # 行内代码
                text = re.sub(r'`(.*?)`', r'\1', part)
                run.text = text
                run.font.name = 'Consolas'
                run.font.size = Pt(9)
            else:
                # 普通文本
                run.text = part
    
    def _format_header_cell(self, cell):
        """格式化表头单元格"""
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # 设置表头字体样式
        for run in paragraph.runs:
            run.font.bold = True
            run.font.size = Pt(11)
        
        # 设置单元格背景色（浅灰色）
        cell._element.get_or_add_tcPr().append(
            self._create_cell_color_element("D9D9D9")
        )
    
    def _format_data_cell(self, cell):
        """格式化数据单元格"""
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        
        # 设置数据字体样式
        for run in paragraph.runs:
            run.font.size = Pt(10)
    
    def _create_cell_color_element(self, color_hex):
        """创建单元格颜色元素"""
        shading_elm = OxmlElement('w:shd')
        shading_elm.set(qn('w:fill'), color_hex)
        return shading_elm
    
    def _set_table_borders(self, table):
        """设置表格边框"""
        tbl = table._tbl
        tblPr = tbl.tblPr
        
        # 创建边框元素
        tblBorders = OxmlElement('w:tblBorders')
        
        # 定义边框样式
        border_attrs = {
            'w:val': 'single',
            'w:sz': '4',
            'w:space': '0',
            'w:color': '000000'
        }
        
        # 添加各种边框
        for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
            border = OxmlElement(f'w:{border_name}')
            for attr, value in border_attrs.items():
                border.set(qn(attr), value)
            tblBorders.append(border)
        
        tblPr.append(tblBorders)


class EnhancedMarkdownToWord:
    """增强的Markdown到Word转换器"""
    
    def __init__(self):
        self.table_converter = AdvancedTableConverter()
    
    def convert(self, markdown_file, output_file):
        """转换Markdown文件到Word"""
        # 读取Markdown文件
        with open(markdown_file, 'r', encoding='utf-8') as f:
            markdown_content = f.read()
        
        # 创建Word文档
        doc = Document()
        
        # 解析表格
        tables = self.table_converter.parse_markdown_table(markdown_content)
        
        # 按行处理内容
        lines = markdown_content.split('\n')
        
        # 跟踪已处理的表格
        processed_tables = set()
        
        i = 0
        while i < len(lines):
            line = lines[i].strip()
            
            # 检查当前行是否在表格中
            current_table = None
            for table_idx, table in enumerate(tables):
                if (table['start_line'] <= i <= table['end_line'] and 
                    table_idx not in processed_tables):
                    current_table = (table_idx, table)
                    break
            
            if current_table:
                # 处理表格
                table_idx, table_data = current_table
                self.table_converter.add_table_to_document(doc, table_data)
                processed_tables.add(table_idx)
                # 跳过表格的所有行
                i = table_data['end_line'] + 1
                continue
            
            # 处理非表格内容
            if line:
                if line.startswith('# '):
                    doc.add_heading(line[2:], level=1)
                elif line.startswith('## '):
                    doc.add_heading(line[3:], level=2)
                elif line.startswith('### '):
                    doc.add_heading(line[4:], level=3)
                elif line.startswith('#### '):
                    doc.add_heading(line[5:], level=4)
                elif line.startswith('- ') or line.startswith('* '):
                    doc.add_paragraph(line[2:], style='List Bullet')
                elif re.match(r'^\d+\.\s', line):
                    doc.add_paragraph(line[3:], style='List Number')
                elif line.startswith('> '):
                    doc.add_paragraph(line[2:], style='Quote')
                elif line.startswith('```'):
                    # 处理代码块
                    i += 1
                    code_content = []
                    while i < len(lines) and not lines[i].strip().startswith('```'):
                        code_content.append(lines[i])
                        i += 1
                    
                    if code_content:
                        code_para = doc.add_paragraph()
                        code_run = code_para.add_run('\n'.join(code_content))
                        code_run.font.name = 'Consolas'
                        code_run.font.size = Pt(9)
                        # 设置代码块背景色
                        code_para.style = 'Normal'
                else:
                    # 普通段落
                    doc.add_paragraph(line)
            else:
                # 空行
                if i < len(lines) - 1:  # 不在文件末尾添加空段落
                    doc.add_paragraph()
            
            i += 1
        
        # 保存文档
        doc.save(output_file)
        print(f"✓ 成功转换文档（含高级表格支持）: {output_file}")
        return True


def main():
    """主函数"""
    converter = EnhancedMarkdownToWord()
    
    # 测试不同的文档
    test_files = [
        ('sample.md', 'enhanced_sample_output.docx'),
        ('table_test.md', 'enhanced_table_test_output.docx')
    ]
    
    for input_file, output_file in test_files:
        try:
            print(f"正在转换 {input_file}...")
            result = converter.convert(input_file, output_file)
            if result:
                print(f"✓ {input_file} 转换成功！表格已正确格式化。")
            else:
                print(f"✗ {input_file} 转换失败")
        except Exception as e:
            print(f"转换 {input_file} 出错: {e}")
            import traceback
            traceback.print_exc()


if __name__ == "__main__":
    main()