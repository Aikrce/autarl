#!/usr/bin/env python3
"""
Table-Enhanced Simple Converter - 集成表格增强功能的简易转换器
"""

import re
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import markdown2


def parse_markdown_tables(markdown_text):
    """解析Markdown中的表格"""
    tables = []
    lines = markdown_text.split('\n')
    
    i = 0
    while i < len(lines):
        line = lines[i].strip()
        # 检测表格开始
        if '|' in line and not line.startswith('```'):
            table_data = []
            table_start = i
            
            # 读取表格的所有行
            while i < len(lines) and '|' in lines[i].strip():
                table_line = lines[i].strip()
                if table_line:
                    # 跳过分隔行
                    if not re.match(r'^\|[\s\-\|:]+\|?$', table_line):
                        # 解析表格行
                        cells = [cell.strip() for cell in table_line.split('|')]
                        # 移除首尾的空元素
                        if cells and cells[0] == '':
                            cells = cells[1:]
                        if cells and cells[-1] == '':
                            cells = cells[:-1]
                        
                        if cells:
                            table_data.append(cells)
                i += 1
            
            if table_data:
                tables.append({
                    'data': table_data,
                    'start_line': table_start,
                    'end_line': i - 1
                })
            i -= 1
        i += 1
        
    return tables


def format_cell_content(cell, content):
    """格式化单元格内容，支持粗体、斜体、代码等"""
    paragraph = cell.paragraphs[0]
    paragraph.clear()
    
    # 处理不同格式的模式
    patterns = [
        (r'\*\*(.*?)\*\*', 'bold'),      # **粗体**
        (r'__(.*?)__', 'bold'),          # __粗体__
        (r'\*(.*?)\*', 'italic'),        # *斜体*
        (r'(?<!_)_(.*?)_(?!_)', 'italic'), # _斜体_
        (r'`(.*?)`', 'code'),            # `代码`
    ]
    
    # 分割内容并应用格式
    parts = []
    current_pos = 0
    
    while current_pos < len(content):
        next_match = None
        next_type = None
        next_pos = len(content)
        
        # 找到最近的格式匹配
        for pattern, format_type in patterns:
            match = re.search(pattern, content[current_pos:])
            if match and match.start() + current_pos < next_pos:
                next_match = match
                next_type = format_type
                next_pos = match.start() + current_pos
        
        # 添加普通文本
        if next_pos > current_pos:
            parts.append(('text', content[current_pos:next_pos]))
        
        # 添加格式化文本
        if next_match:
            parts.append((next_type, next_match.group(1)))
            current_pos = next_pos + len(next_match.group(0))
        else:
            break
    
    # 生成格式化的段落
    for part_type, text in parts:
        if not text:
            continue
            
        run = paragraph.add_run(text)
        
        if part_type == 'bold':
            run.bold = True
        elif part_type == 'italic':
            run.italic = True
        elif part_type == 'code':
            run.font.name = 'Consolas'
            run.font.size = Pt(9)


def add_table_to_doc(doc, table_data):
    """将表格添加到Word文档"""
    if not table_data or not table_data['data']:
        return
        
    rows = table_data['data']
    if not rows:
        return
        
    # 确定表格尺寸
    max_cols = max(len(row) for row in rows)
    if max_cols == 0:
        return
        
    # 创建表格
    table = doc.add_table(rows=len(rows), cols=max_cols)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # 填充表格内容
    for row_idx, row_data in enumerate(rows):
        for col_idx in range(max_cols):
            cell = table.cell(row_idx, col_idx)
            
            # 获取单元格内容
            if col_idx < len(row_data):
                cell_content = row_data[col_idx].strip()
            else:
                cell_content = ""
            
            # 设置内容和格式
            if cell_content:
                format_cell_content(cell, cell_content)
            
            # 设置表头样式
            if row_idx == 0:
                paragraph = cell.paragraphs[0]
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.size = Pt(11)
                
                # 设置表头背景色
                shading_elm = OxmlElement('w:shd')
                shading_elm.set(qn('w:fill'), "D9D9D9")
                cell._element.get_or_add_tcPr().append(shading_elm)
            else:
                # 数据行样式
                for run in cell.paragraphs[0].runs:
                    run.font.size = Pt(10)
    
    # 添加表格后的空行
    doc.add_paragraph()


def enhanced_markdown_to_docx(markdown_file, output_file):
    """增强的Markdown到Word转换器，支持表格"""
    
    # 读取Markdown文件
    with open(markdown_file, 'r', encoding='utf-8') as f:
        markdown_content = f.read()
    
    # 创建Word文档
    doc = Document()
    
    # 解析表格
    tables = parse_markdown_tables(markdown_content)
    
    # 按行处理内容
    lines = markdown_content.split('\n')
    processed_tables = set()
    
    i = 0
    while i < len(lines):
        line = lines[i].strip()
        
        # 检查是否在表格中
        current_table = None
        for table_idx, table in enumerate(tables):
            if (table['start_line'] <= i <= table['end_line'] and 
                table_idx not in processed_tables):
                current_table = (table_idx, table)
                break
        
        if current_table:
            # 处理表格
            table_idx, table_data = current_table
            add_table_to_doc(doc, table_data)
            processed_tables.add(table_idx)
            i = table_data['end_line'] + 1
            continue
        
        # 处理其他内容
        if line:
            if line.startswith('# '):
                doc.add_heading(line[2:], level=1)
            elif line.startswith('## '):
                doc.add_heading(line[3:], level=2)
            elif line.startswith('### '):
                doc.add_heading(line[4:], level=3)
            elif line.startswith('#### '):
                doc.add_heading(line[5:], level=4)
            elif line.startswith('- ') or line.startswith('* '):
                doc.add_paragraph(line[2:], style='List Bullet')
            elif re.match(r'^\d+\.\s', line):
                doc.add_paragraph(line[3:], style='List Number')
            elif line.startswith('> '):
                doc.add_paragraph(line[2:], style='Quote')
            elif line.startswith('```'):
                # 代码块处理
                i += 1
                code_content = []
                while i < len(lines) and not lines[i].strip().startswith('```'):
                    code_content.append(lines[i])
                    i += 1
                
                if code_content:
                    code_para = doc.add_paragraph()
                    code_run = code_para.add_run('\n'.join(code_content))
                    code_run.font.name = 'Consolas'
                    code_run.font.size = Pt(9)
            else:
                # 普通段落
                doc.add_paragraph(line)
        else:
            # 空行
            if i < len(lines) - 1:
                doc.add_paragraph()
        
        i += 1
    
    # 保存文档
    doc.save(output_file)
    print(f"✓ 转换完成（支持高级表格）: {output_file}")
    return True


if __name__ == "__main__":
    try:
        # 转换示例文件
        result1 = enhanced_markdown_to_docx('sample.md', 'sample_with_tables.docx')
        # 转换表格测试文件
        result2 = enhanced_markdown_to_docx('table_test.md', 'table_test_output.docx')
        
        if result1 and result2:
            print("\n🎉 所有转换成功！")
            print("生成的文件包含：")
            print("- sample_with_tables.docx（基础示例）")
            print("- table_test_output.docx（表格测试）")
            print("\n表格特性：")
            print("✅ 表头自动加粗并居中")
            print("✅ 表头背景色区分")
            print("✅ 支持粗体、斜体、代码格式")
            print("✅ 自动处理不规则表格")
            print("✅ 专业的表格边框样式")
        else:
            print("转换失败")
    except Exception as e:
        print(f"转换出错: {e}")
        import traceback
        traceback.print_exc()