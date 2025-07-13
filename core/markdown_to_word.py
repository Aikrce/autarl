#!/usr/bin/env python3

import os
import argparse
from pathlib import Path
import pypandoc
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
import markdown2
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from templates_config import get_template, list_templates
from document_analyzer import analyze_markdown_document
import logging

logger = logging.getLogger(__name__)


class MarkdownToWordConverter:
    def __init__(self, template_name='default'):
        self.doc = None
        self.template_name = template_name
        self.template = get_template(template_name)
        self.content_analysis = None
        self.smart_matching = True  # 启用智能模板匹配
        
    def convert_with_pandoc(self, input_file, output_file, output_format='docx'):
        """使用pandoc进行转换（推荐方式）"""
        try:
            extra_args = ['--standalone']
            if output_format == 'docx':
                extra_args.append('--toc')
            elif output_format == 'pdf':
                extra_args.extend(['--pdf-engine=xelatex', '--toc'])
            elif output_format == 'html':
                extra_args.extend(['--toc', '--css=style.css'])
            
            pypandoc.convert_file(
                input_file,
                output_format,
                outputfile=output_file,
                extra_args=extra_args
            )
            print(f"✓ 成功转换: {input_file} -> {output_file}")
            return True
        except Exception as e:
            print(f"✗ Pandoc转换失败: {e}")
            return False
    
    def convert_to_html(self, input_file, output_file):
        """转换为HTML格式"""
        try:
            with open(input_file, 'r', encoding='utf-8') as f:
                markdown_text = f.read()
            
            html = markdown2.markdown(
                markdown_text,
                extras=['tables', 'fenced-code-blocks', 'header-ids', 'toc']
            )
            
            html_template = f"""<!DOCTYPE html>
<html lang="zh-CN">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Markdown Document</title>
    <style>
        body {{
            font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', 'Microsoft YaHei', sans-serif;
            line-height: 1.6;
            color: #333;
            max-width: 800px;
            margin: 0 auto;
            padding: 20px;
        }}
        code {{
            background-color: #f4f4f4;
            padding: 2px 4px;
            border-radius: 3px;
        }}
        pre {{
            background-color: #f4f4f4;
            padding: 10px;
            border-radius: 5px;
            overflow-x: auto;
        }}
        blockquote {{
            border-left: 4px solid #ddd;
            padding-left: 16px;
            margin-left: 0;
        }}
        table {{
            border-collapse: collapse;
            width: 100%;
        }}
        th, td {{
            border: 1px solid #ddd;
            padding: 8px;
            text-align: left;
        }}
        th {{
            background-color: #f2f2f2;
        }}
    </style>
</head>
<body>
{html}
</body>
</html>"""
            
            with open(output_file, 'w', encoding='utf-8') as f:
                f.write(html_template)
            
            print(f"✓ 成功转换: {input_file} -> {output_file}")
            return True
            
        except Exception as e:
            print(f"✗ HTML转换失败: {e}")
            return False
    
    def convert_to_txt(self, input_file, output_file):
        """转换为纯文本格式"""
        try:
            with open(input_file, 'r', encoding='utf-8') as f:
                markdown_text = f.read()
            
            # 简单的Markdown到纯文本转换
            import re
            
            # 移除图片链接
            text = re.sub(r'!\[.*?\]\(.*?\)', '', markdown_text)
            # 移除链接，保留链接文本
            text = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', text)
            # 移除代码块标记
            text = re.sub(r'```[^`]*```', '', text, flags=re.DOTALL)
            # 移除行内代码标记
            text = re.sub(r'`([^`]+)`', r'\1', text)
            # 移除粗体和斜体标记
            text = re.sub(r'\*\*([^*]+)\*\*', r'\1', text)
            text = re.sub(r'__([^_]+)__', r'\1', text)
            text = re.sub(r'\*([^*]+)\*', r'\1', text)
            text = re.sub(r'_([^_]+)_', r'\1', text)
            # 移除标题标记
            text = re.sub(r'^#+\s*', '', text, flags=re.MULTILINE)
            # 移除列表标记
            text = re.sub(r'^[\s]*[-*+]\s*', '', text, flags=re.MULTILINE)
            text = re.sub(r'^\s*\d+\.\s*', '', text, flags=re.MULTILINE)
            
            with open(output_file, 'w', encoding='utf-8') as f:
                f.write(text)
            
            print(f"✓ 成功转换: {input_file} -> {output_file}")
            return True
            
        except Exception as e:
            print(f"✗ TXT转换失败: {e}")
            return False
    
    def convert_with_python_docx(self, input_file, output_file):
        """使用python-docx进行转换（备用方式）"""
        try:
            with open(input_file, 'r', encoding='utf-8') as f:
                markdown_text = f.read()
            
            # 智能文档分析
            if self.smart_matching:
                logger.info("开始智能文档分析...")
                self.content_analysis = analyze_markdown_document(markdown_text)
                self._log_analysis_results()
            
            # 使用markdown2解析Markdown
            html = markdown2.markdown(
                markdown_text,
                extras=['tables', 'fenced-code-blocks', 'header-ids']
            )
            
            # 创建Word文档
            self.doc = Document()
            self._setup_styles()
            
            # 智能模板匹配：只对存在的内容应用模板格式
            if self.smart_matching and self.content_analysis:
                self._apply_smart_template_matching(markdown_text)
            else:
                # 传统方式：解析并添加所有内容
                self._parse_markdown_content(markdown_text)
            
            # 设置页码系统
            self._setup_page_numbering()
            
            # 保存文档
            self.doc.save(output_file)
            print(f"✓ 成功转换: {input_file} -> {output_file}")
            
            if self.smart_matching and self.content_analysis:
                self._print_matching_summary()
            
            return True
            
        except Exception as e:
            logger.error(f"Python-docx转换失败: {e}")
            print(f"✗ Python-docx转换失败: {e}")
            return False
    
    def _setup_styles(self):
        """设置文档样式"""
        # 应用选定的模板
        self.template.apply_to_document(self.doc)
    
    def _parse_markdown_content(self, content):
        """解析Markdown内容并添加到Word文档"""
        lines = content.split('\n')
        current_paragraph = None
        in_code_block = False
        code_content = []
        
        for line in lines:
            # 代码块处理
            if line.strip().startswith('```'):
                if in_code_block:
                    # 结束代码块
                    code_text = '\n'.join(code_content)
                    p = self.doc.add_paragraph()
                    run = p.add_run(code_text)
                    run.font.name = 'Consolas'
                    run.font.size = Pt(10)
                    run.font.color.rgb = RGBColor(0, 0, 0)
                    # 设置段落样式
                    p.paragraph_format.left_indent = Pt(18)
                    self._add_border(p)
                    code_content = []
                    in_code_block = False
                else:
                    # 开始代码块
                    in_code_block = True
                continue
            
            if in_code_block:
                code_content.append(line)
                continue
            
            # 标题处理 - 为东北师大模板特别优化
            if line.startswith('#'):
                level = len(line) - len(line.lstrip('#'))
                if level <= 6:
                    title_text = line.lstrip('#').strip()
                    
                    if self.template_name == 'nenu_thesis':
                        # 东北师大模板专用标题处理
                        if level == 1:
                            # 一级标题：三号黑体，居中，段前48磅，段后24磅
                            heading_para = self.doc.add_heading(title_text, level=1)
                        elif level == 2:
                            # 二级标题：四号黑体，两端对齐，段前6磅，段后0磅
                            heading_para = self.doc.add_heading(title_text, level=2)
                        elif level == 3:
                            # 三级标题：小四号宋体加粗，两端对齐，段前6磅，段后0磅
                            heading_para = self.doc.add_heading(title_text, level=3)
                        elif level == 4:
                            # 四级标题：小四号宋体，两端对齐，段前0行，段后0行
                            heading_para = self.doc.add_heading(title_text, level=4)
                        else:
                            # 更深层级标题使用默认格式
                            heading_para = self.doc.add_heading(title_text, level=level)
                    else:
                        # 其他模板使用默认标题处理
                        self.doc.add_heading(title_text, level=level)
                    continue
            
            # 列表处理
            if line.strip().startswith(('- ', '* ', '+ ')):
                list_text = line.strip()[2:]
                p = self.doc.add_paragraph(list_text, style='List Bullet')
                continue
            
            if line.strip() and line.strip()[0].isdigit() and line.strip().find('.') == 1:
                list_text = line.strip()[3:]
                p = self.doc.add_paragraph(list_text, style='List Number')
                continue
            
            # 分隔线处理
            if line.strip() in ['---', '***', '___']:
                # 添加分页符
                self.doc.add_page_break()
                continue
            
            # 空行处理
            if not line.strip():
                if current_paragraph:
                    current_paragraph = None
                continue
            
            # 普通段落
            if current_paragraph is None:
                current_paragraph = self.doc.add_paragraph()
            
            # 处理行内格式
            self._process_inline_formatting(current_paragraph, line)
    
    def _process_inline_formatting(self, paragraph, text):
        """处理行内格式（粗体、斜体、代码等）"""
        import re
        
        # 简单的格式处理
        parts = re.split(r'(\*\*.*?\*\*|__.*?__|_.*?_|\*.*?\*|`.*?`)', text)
        
        for part in parts:
            if not part:
                continue
                
            # 粗体
            if part.startswith('**') and part.endswith('**'):
                run = paragraph.add_run(part[2:-2])
                run.bold = True
            elif part.startswith('__') and part.endswith('__'):
                run = paragraph.add_run(part[2:-2])
                run.bold = True
            # 斜体
            elif part.startswith('*') and part.endswith('*') and len(part) > 2:
                run = paragraph.add_run(part[1:-1])
                run.italic = True
            elif part.startswith('_') and part.endswith('_') and len(part) > 2:
                run = paragraph.add_run(part[1:-1])
                run.italic = True
            # 行内代码
            elif part.startswith('`') and part.endswith('`'):
                run = paragraph.add_run(part[1:-1])
                run.font.name = 'Consolas'
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(219, 48, 105)
            else:
                paragraph.add_run(part)
    
    def _add_border(self, paragraph):
        """为段落添加边框"""
        pPr = paragraph._p.get_or_add_pPr()
        borders = OxmlElement('w:pBdr')
        pPr.insert_element_before(borders,
            'w:shd', 'w:tabs', 'w:suppressAutoHyphens', 'w:kinsoku',
            'w:wordWrap', 'w:overflowPunct', 'w:topLinePunct', 'w:autoSpaceDE',
            'w:autoSpaceDN', 'w:bidi', 'w:adjustRightInd', 'w:snapToGrid',
            'w:spacing', 'w:ind', 'w:contextualSpacing', 'w:mirrorIndents',
            'w:suppressOverlap', 'w:jc', 'w:textDirection', 'w:textAlignment',
            'w:textboxTightWrap', 'w:outlineLvl', 'w:divId', 'w:cnfStyle',
            'w:rPr', 'w:sectPr', 'w:pPrChange'
        )
        
        for border_name in ['top', 'left', 'bottom', 'right']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), '4')
            border.set(qn('w:space'), '1')
            border.set(qn('w:color'), 'auto')
            borders.append(border)
    
    def _log_analysis_results(self):
        """记录文档分析结果"""
        if not self.content_analysis:
            return
        
        detected = self.content_analysis['detected_components']
        doc_type = self.content_analysis['document_type']
        sections_count = len(self.content_analysis['sections'])
        
        logger.info(f"文档类型: {doc_type}")
        logger.info(f"检测到章节: {sections_count}")
        logger.info(f"检测到学术组件: {', '.join(detected) if detected else '无'}")
    
    def _print_matching_summary(self):
        """打印智能匹配摘要"""
        if not self.content_analysis:
            return
            
        detected = self.content_analysis['detected_components']
        missing = self._get_missing_components()
        
        print(f"\n📊 智能模板匹配摘要:")
        print(f"   文档类型: {self.content_analysis['document_type']}")
        print(f"   已匹配组件: {len(detected)} 个 ({', '.join(detected) if detected else '无'})")
        print(f"   保留模板组件: {len(missing)} 个 ({', '.join(missing) if missing else '无'})")
    
    def _get_missing_components(self):
        """获取模板中缺失的组件"""
        if not self.content_analysis:
            return set()
        
        # 检查是否有分析器的get_missing_components方法
        try:
            from document_analyzer import MarkdownDocumentAnalyzer
            analyzer = MarkdownDocumentAnalyzer()
            return analyzer.get_missing_components(
                self.template_name, 
                self.content_analysis['detected_components']
            )
        except:
            return set()
    
    def _apply_smart_template_matching(self, markdown_text):
        """智能模板匹配：按正确顺序排列内容"""
        detected_components = self.content_analysis['detected_components']
        content_mapping = self.content_analysis['content_mapping']
        sections = self.content_analysis['sections']
        
        # 定义正确的章节顺序
        section_order = [
            'cover_page',
            'english_cover', 
            'declaration',
            'authorization',
            'abstract_cn',
            'abstract_en', 
            'toc',
            'symbols',
            'figures_list',
            'tables_list',
            'introduction',
            'literature_review', 
            'methodology',
            'results',
            'discussion',
            'conclusion',
            'references',
            'appendix',
            'acknowledgments'
        ]
        
        # 按正确顺序添加检测到的内容
        sections_by_type = {}
        for section in sections:
            if section.section_type not in sections_by_type:
                sections_by_type[section.section_type] = []
            sections_by_type[section.section_type].append(section)
        
        # 按顺序处理各个部分
        for section_type in section_order:
            if section_type in sections_by_type:
                for section in sections_by_type[section_type]:
                    self._add_academic_component(section.section_type, section.content, section.name)
            elif section_type in ['cover_page', 'english_cover', 'declaration', 'authorization']:
                # 添加缺失的必要组件
                missing_components = self._get_missing_components()
                if section_type in missing_components:
                    self._add_template_placeholders({section_type})
        
        # 处理未分类的章节内容
        if 'unknown' in sections_by_type:
            for section in sections_by_type['unknown']:
                self._add_regular_content(section.content)
        
        # 为缺失的其他组件添加模板占位符
        missing_components = self._get_missing_components()
        remaining_missing = missing_components - {'cover_page', 'english_cover', 'declaration', 'authorization'}
        if remaining_missing and self.template_name == 'nenu_thesis':
            self._add_template_placeholders(remaining_missing)
    
    def _add_detected_content(self, sections, content_mapping):
        """添加检测到的内容"""
        for section in sections:
            if section.section_type == 'unknown':
                # 未识别的内容按正文处理
                self._add_regular_content(section.content)
            else:
                # 按照学术组件类型格式化
                self._add_academic_component(section.section_type, section.content, section.name)
    
    def _add_regular_content(self, content):
        """添加常规内容（按正文格式）"""
        lines = content.split('\n')
        current_paragraph = None
        in_code_block = False
        code_content = []
        
        for line in lines:
            # 代码块处理
            if line.strip().startswith('```'):
                if in_code_block:
                    # 结束代码块
                    code_text = '\n'.join(code_content)
                    p = self.doc.add_paragraph()
                    run = p.add_run(code_text)
                    run.font.name = 'Consolas'
                    run.font.size = Pt(10)
                    run.font.color.rgb = RGBColor(0, 0, 0)
                    # 设置段落样式
                    p.paragraph_format.left_indent = Pt(18)
                    self._add_border(p)
                    code_content = []
                    in_code_block = False
                else:
                    # 开始代码块
                    in_code_block = True
                continue
            
            if in_code_block:
                code_content.append(line)
                continue
            
            # 标题处理
            if line.startswith('#'):
                level = len(line) - len(line.lstrip('#'))
                if level <= 6:
                    title_text = line.lstrip('#').strip()
                    self.doc.add_heading(title_text, level=level)
                    continue
            
            # 列表处理
            if line.strip().startswith(('- ', '* ', '+ ')):
                list_text = line.strip()[2:]
                p = self.doc.add_paragraph(list_text, style='List Bullet')
                continue
            
            if line.strip() and line.strip()[0].isdigit() and line.strip().find('.') == 1:
                list_text = line.strip()[3:]
                p = self.doc.add_paragraph(list_text, style='List Number')
                continue
            
            # 分隔线处理
            if line.strip() in ['---', '***', '___']:
                # 添加分页符
                self.doc.add_page_break()
                continue
            
            # 空行处理
            if not line.strip():
                if current_paragraph:
                    current_paragraph = None
                continue
            
            # 普通段落
            if current_paragraph is None:
                current_paragraph = self.doc.add_paragraph()
            
            # 处理行内格式
            self._process_inline_formatting(current_paragraph, line)
    
    def _add_academic_component(self, component_type, content, section_name):
        """根据学术组件类型添加格式化内容"""
        if component_type in ['abstract_cn', 'abstract_en']:
            self._add_abstract_section(component_type, content, section_name)
        elif component_type in ['keywords_cn', 'keywords_en']:
            self._add_keywords_section(component_type, content)
        elif component_type == 'references':
            self._add_references_section(content, section_name)
        elif component_type == 'toc':
            self._add_toc_section(content, section_name)
        elif component_type == 'appendix':
            self._add_appendix_section(content, section_name)
        elif component_type in ['introduction', 'conclusion', 'methodology', 'results', 'discussion']:
            self._add_chapter_section(content, section_name)
        else:
            # 其他类型按常规内容处理
            self._add_regular_content(content)
    
    def _add_abstract_section(self, component_type, content, section_name):
        """添加摘要部分 - 精确按照模板要求格式化"""
        # 添加分页符
        self.doc.add_page_break()
        
        # 添加摘要标题：正确实现中间空2个汉字
        if component_type == 'abstract_cn':
            # 中文摘要标题："摘　　要" - 使用全角空格实现精确间距
            title_text = '摘　　要'  # 使用全角空格实现间距
            style_name = 'Abstract Title CN'
        else:
            # 英文摘要标题
            title_text = 'Abstract'
            style_name = 'Abstract Title EN'
        
        # 检查样式是否存在并创建标题
        try:
            title_para = self.doc.add_paragraph(title_text, style=style_name)
        except:
            title_para = self.doc.add_paragraph(title_text)
            # 手动设置格式：三号黑体，居中，段前48磅，段后24磅，1.5倍行距
            title_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            for run in title_para.runs:
                if component_type == 'abstract_cn':
                    run.font.name = '黑体'
                else:
                    run.font.name = 'Times New Roman'
                run.font.size = Pt(16)  # 三号
                run.bold = True
                run.font.color.rgb = RGBColor(0, 0, 0)  # 确保黑色
            title_para.paragraph_format.space_before = Pt(48)
            title_para.paragraph_format.space_after = Pt(24)
            title_para.paragraph_format.line_spacing = 1.5
        
        # 添加摘要内容
        lines = content.split('\n')[1:]  # 跳过标题行
        
        # 过滤掉空行和只包含空格的行，处理实际的摘要内容
        content_lines = [line.strip() for line in lines if line.strip() and not line.strip().startswith('关键词') and not line.strip().startswith('Key words')]
        
        # 添加摘要正文段落
        for line in content_lines:
            if line.strip():
                try:
                    if component_type == 'abstract_cn':
                        para = self.doc.add_paragraph(line, style='Abstract Body CN')
                    else:
                        para = self.doc.add_paragraph(line, style='Abstract Body EN')
                except:
                    para = self.doc.add_paragraph(line)
                    # 手动设置摘要正文格式：小四号，两端对齐，首行缩进2字符，1.5倍行距
                    for run in para.runs:
                        if component_type == 'abstract_cn':
                            run.font.name = '宋体'
                        else:
                            run.font.name = 'Times New Roman'
                        run.font.size = Pt(12)  # 小四号
                        run.font.color.rgb = RGBColor(0, 0, 0)  # 确保黑色
                    para.paragraph_format.first_line_indent = Cm(0.74)  # 首行缩进2字符
                    para.paragraph_format.line_spacing = 1.5
                    para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        
        # 添加空行（在关键词前）
        self.doc.add_paragraph("")
        
        # 查找并添加关键词
        for line in lines:
            if line.strip():
                if (component_type == 'abstract_cn' and '关键词' in line) or \
                   (component_type == 'abstract_en' and 'Key words' in line):
                    try:
                        if component_type == 'abstract_cn':
                            kw_para = self.doc.add_paragraph(line, style='Keywords CN')
                        else:
                            kw_para = self.doc.add_paragraph(line, style='Keywords EN')
                    except:
                        kw_para = self.doc.add_paragraph(line)
                        # 手动设置关键词格式
                        for run in kw_para.runs:
                            if component_type == 'abstract_cn':
                                run.font.name = '宋体'
                                run.bold = True  # 关键词标签加粗
                            else:
                                run.font.name = 'Times New Roman'
                                run.bold = True
                            run.font.size = Pt(12)  # 小四号
                            run.font.color.rgb = RGBColor(0, 0, 0)  # 确保黑色
                        
                        if component_type == 'abstract_cn':
                            kw_para.paragraph_format.first_line_indent = Cm(0.74)  # 首行缩进2字符
                            kw_para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                        else:
                            # 英文关键词：悬挂缩进5.95字符（约2.1cm）
                            kw_para.paragraph_format.hanging_indent = Cm(2.1)
                            kw_para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                        
                        kw_para.paragraph_format.line_spacing = 1.5
                    break  # 只处理第一个关键词行
    
    def _add_keywords_section(self, component_type, content):
        """添加关键词部分"""
        try:
            if component_type == 'keywords_cn':
                para = self.doc.add_paragraph(content, style='Keywords CN')
            else:
                para = self.doc.add_paragraph(content, style='Keywords EN')
        except:
            para = self.doc.add_paragraph(content)
    
    def _add_references_section(self, content, section_name):
        """添加参考文献部分"""
        # 添加参考文献标题
        try:
            title_para = self.doc.add_paragraph(section_name, style='Reference Title')
        except:
            title_para = self.doc.add_paragraph(section_name)
            title_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            for run in title_para.runs:
                run.font.name = '黑体'
                run.font.size = Pt(16)  # 三号
                run.bold = True
                run.font.color.rgb = RGBColor(0, 0, 0)  # 确保黑色
            title_para.paragraph_format.space_before = Pt(48)
            title_para.paragraph_format.space_after = Pt(24)
        
        # 添加参考文献内容，并设置智能悬挂缩进
        lines = content.split('\n')[1:]  # 跳过标题行
        for line in lines:
            if line.strip():
                # 提取编号以决定悬挂缩进
                import re
                number_match = re.match(r'^\[(\d+)\]', line.strip())
                
                try:
                    para = self.doc.add_paragraph(line, style='Reference Content')
                except:
                    para = self.doc.add_paragraph(line)
                    # 手动设置格式
                    for run in para.runs:
                        run.font.name = '宋体'
                        run.font.size = Pt(12)  # 小四号
                        run.font.color.rgb = RGBColor(0, 0, 0)  # 确保黑色
                    para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                    para.paragraph_format.line_spacing = 1.0
                
                # 根据编号设置悬挂缩进
                if number_match:
                    number = int(number_match.group(1))
                    if number <= 9:
                        para.paragraph_format.hanging_indent = Cm(0.6)  # 1-9编号
                    elif number <= 99:
                        para.paragraph_format.hanging_indent = Cm(0.74)  # 10-99编号
                    else:
                        para.paragraph_format.hanging_indent = Cm(0.9)  # 100+编号
                else:
                    # 默认悬挂缩进
                    para.paragraph_format.hanging_indent = Cm(0.6)
    
    def _add_toc_section(self, content, section_name):
        """添加目录部分 - 精确按照模板要求格式化"""
        # 添加分页符
        self.doc.add_page_break()
        
        # 添加目录标题 - "目　　录" (中间空2个汉字)
        title_text = '目　　录'  # 使用全角空格实现间距
        try:
            title_para = self.doc.add_paragraph(title_text, style='TOC Title')
        except:
            title_para = self.doc.add_paragraph(title_text)
            title_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            # 设置目录标题格式：三号黑体，居中，段前48磅，段后24磅，1.5倍行距
            for run in title_para.runs:
                run.font.name = '黑体'
                run.font.size = Pt(16)  # 三号
                run.bold = True
                run.font.color.rgb = RGBColor(0, 0, 0)  # 确保黑色
            title_para.paragraph_format.space_before = Pt(48)
            title_para.paragraph_format.space_after = Pt(24)
            title_para.paragraph_format.line_spacing = 1.5
        
        # 解析目录内容并格式化
        lines = content.split('\n')[1:]  # 跳过标题行
        for line in lines:
            line = line.strip()
            if not line:
                continue
            
            # 清理markdown语法符号
            cleaned_line = line
            # 移除所有markdown标记
            cleaned_line = cleaned_line.replace('**', '')  # 移除粗体标记
            cleaned_line = cleaned_line.replace('*', '')   # 移除斜体标记
            cleaned_line = cleaned_line.replace('_', '')   # 移除下划线标记
            cleaned_line = cleaned_line.strip()
            
            if not cleaned_line:
                continue
                
            # 判断目录层级和类型
            if ('第' in cleaned_line and '章' in cleaned_line) or cleaned_line in ['绪论', '结论', '引言']:
                # 章标题 - 黑体小四号，首行无缩进
                try:
                    para = self.doc.add_paragraph(cleaned_line, style='TOC Level 1')
                except:
                    para = self.doc.add_paragraph(cleaned_line)
                    for run in para.runs:
                        run.font.name = '黑体'
                        run.font.size = Pt(12)  # 小四号
                        run.font.color.rgb = RGBColor(0, 0, 0)  # 黑色
                    para.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                    para.paragraph_format.left_indent = Pt(0)
                    para.paragraph_format.line_spacing = 1.5
                    para.paragraph_format.space_before = Pt(0)
                    para.paragraph_format.space_after = Pt(0)
                    
            elif '.' in cleaned_line and any(char.isdigit() for char in cleaned_line.split('.')[0]):
                # 二级目录项（如"1.1 相关理论基础"）- 宋体小四号，左缩进1字符
                try:
                    para = self.doc.add_paragraph(cleaned_line, style='TOC Level 2')
                except:
                    para = self.doc.add_paragraph(cleaned_line)
                    for run in para.runs:
                        run.font.name = '宋体'
                        run.font.size = Pt(12)  # 小四号
                        run.font.color.rgb = RGBColor(0, 0, 0)  # 黑色
                    para.paragraph_format.left_indent = Cm(0.37)  # 左缩进1字符
                    para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                    para.paragraph_format.line_spacing = 1.5
                    para.paragraph_format.space_before = Pt(0)
                    para.paragraph_format.space_after = Pt(0)
                    
            elif cleaned_line.count('.') >= 2:
                # 三级目录项（如"1.1.1 具体内容"）- 宋体小四号，左缩进2字符
                try:
                    para = self.doc.add_paragraph(cleaned_line, style='TOC Level 3')
                except:
                    para = self.doc.add_paragraph(cleaned_line)
                    for run in para.runs:
                        run.font.name = '宋体'
                        run.font.size = Pt(12)  # 小四号
                        run.font.color.rgb = RGBColor(0, 0, 0)  # 黑色
                    para.paragraph_format.left_indent = Cm(0.74)  # 左缩进2字符
                    para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                    para.paragraph_format.line_spacing = 1.5
                    para.paragraph_format.space_before = Pt(0)
                    para.paragraph_format.space_after = Pt(0)
                    
            else:
                # 其他目录项 - 默认为一级格式
                try:
                    para = self.doc.add_paragraph(cleaned_line, style='TOC Level 1')
                except:
                    para = self.doc.add_paragraph(cleaned_line)
                    for run in para.runs:
                        run.font.name = '黑体'
                        run.font.size = Pt(12)  # 小四号
                        run.font.color.rgb = RGBColor(0, 0, 0)  # 黑色
                    para.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                    para.paragraph_format.left_indent = Pt(0)
                    para.paragraph_format.line_spacing = 1.5
                    para.paragraph_format.space_before = Pt(0)
                    para.paragraph_format.space_after = Pt(0)
    
    def _add_appendix_section(self, content, section_name):
        """添加附录部分"""
        # 不单独分页，与前面内容连续
        
        # 添加附录标题
        try:
            title_para = self.doc.add_paragraph(section_name, style='Appendix Title')
        except:
            title_para = self.doc.add_paragraph(section_name)
            title_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            for run in title_para.runs:
                run.font.name = '黑体'
                run.font.size = Pt(16)  # 三号
                run.bold = True
                run.font.color.rgb = RGBColor(0, 0, 0)  # 确保黑色
            title_para.paragraph_format.space_before = Pt(48)
            title_para.paragraph_format.space_after = Pt(24)
        
        # 添加附录内容
        lines = content.split('\n')[1:]  # 跳过标题行
        for line in lines:
            if line.strip():
                try:
                    para = self.doc.add_paragraph(line, style='Appendix Content')
                except:
                    para = self.doc.add_paragraph(line)
                    for run in para.runs:
                        run.font.name = '宋体'
                        run.font.size = Pt(12)  # 小四号
                        run.font.color.rgb = RGBColor(0, 0, 0)  # 确保黑色
                    para.paragraph_format.first_line_indent = Cm(0.74)  # 首行缩进2字符
                    para.paragraph_format.line_spacing = 1.5
                    para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    
    def _add_chapter_section(self, content, section_name):
        """添加章节部分 - 精确按照模板要求格式化"""
        # 为主要章节添加分页符
        if any(keyword in section_name for keyword in ['第一章', '第二章', '第三章', '第四章', '第五章', '第六章', '引言', '绪论', '结论']):
            self.doc.add_page_break()
        
        # 添加章节标题（使用适当的格式）
        if self.template_name == 'nenu_thesis':
            # 东北师大格式：章标题居中，三号黑体，加粗，段前48磅，段后24磅，1.5倍行距
            heading_para = self.doc.add_paragraph(section_name)
            heading_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            for run in heading_para.runs:
                run.font.name = '黑体'
                run.font.size = Pt(16)  # 三号
                run.bold = True
                run.font.color.rgb = RGBColor(0, 0, 0)  # 确保黑色
            heading_para.paragraph_format.space_before = Pt(48)
            heading_para.paragraph_format.space_after = Pt(24)
            heading_para.paragraph_format.line_spacing = 1.5
        else:
            # 其他模板使用默认标题样式
            self.doc.add_heading(section_name, level=1)
        
        # 添加章节内容
        content_lines = content.split('\n')[1:]  # 跳过标题行
        
        current_paragraph = None
        for line in content_lines:
            line = line.strip()
            if not line:
                # 空行处理
                if current_paragraph:
                    current_paragraph = None
                continue
            
            # 检测二级标题
            if line.startswith('##'):
                # 二级标题：四号黑体，两端对齐，段前6磅，段后0磅，1.5倍行距
                title_text = line.lstrip('#').strip()
                if self.template_name == 'nenu_thesis':
                    heading_para = self.doc.add_paragraph(title_text)
                    heading_para.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                    for run in heading_para.runs:
                        run.font.name = '黑体'
                        run.font.size = Pt(14)  # 四号
                        run.bold = True
                        run.font.color.rgb = RGBColor(0, 0, 0)  # 确保黑色
                    heading_para.paragraph_format.space_before = Pt(6)
                    heading_para.paragraph_format.space_after = Pt(0)
                    heading_para.paragraph_format.line_spacing = 1.5
                else:
                    self.doc.add_heading(title_text, level=2)
                current_paragraph = None
                continue
            
            # 检测三级标题
            elif line.startswith('###'):
                # 三级标题：小四号宋体加粗，两端对齐，段前6磅，段后0磅，1.5倍行距
                title_text = line.lstrip('#').strip()
                if self.template_name == 'nenu_thesis':
                    heading_para = self.doc.add_paragraph(title_text)
                    heading_para.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                    for run in heading_para.runs:
                        run.font.name = '宋体'
                        run.font.size = Pt(12)  # 小四号
                        run.bold = True
                        run.font.color.rgb = RGBColor(0, 0, 0)  # 确保黑色
                    heading_para.paragraph_format.space_before = Pt(6)
                    heading_para.paragraph_format.space_after = Pt(0)
                    heading_para.paragraph_format.line_spacing = 1.5
                else:
                    self.doc.add_heading(title_text, level=3)
                current_paragraph = None
                continue
            
            # 普通段落文字：中文宋体，英文Times New Roman，小四号，两端对齐，首行缩进2字符，1.5倍行距
            if current_paragraph is None:
                current_paragraph = self.doc.add_paragraph()
                # 设置段落格式
                current_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                current_paragraph.paragraph_format.first_line_indent = Cm(0.74)  # 首行缩进2字符
                current_paragraph.paragraph_format.line_spacing = 1.5
                current_paragraph.paragraph_format.space_before = Pt(0)
                current_paragraph.paragraph_format.space_after = Pt(0)
            
            # 处理行内格式并添加文本
            self._process_inline_formatting(current_paragraph, line)
            
            # 设置字体
            for run in current_paragraph.runs:
                run.font.name = '宋体'  # 中文采用宋体
                run.font.size = Pt(12)  # 小四号
                run.font.color.rgb = RGBColor(0, 0, 0)  # 确保黑色
                # 设置中英文字体
                if hasattr(run, '_element'):
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                    run._element.rPr.rFonts.set(qn('w:ascii'), 'Times New Roman')
                    run._element.rPr.rFonts.set(qn('w:hAnsi'), 'Times New Roman')
    
    def _get_cover_content(self):
        """生成中文封面内容"""
        return """
学校代码：10200                           研究生学号：【学号】
                                        密级：公开

东北师范大学硕士学位论文

【中文论文题目】

                                        
                                        学    位    申    请    人：【姓名】
                                        指    导    教    师：【导师姓名】  【职称】
                                        学    科、专    业：【专业名称】
                                        研    究    方    向：【研究方向】
                                        
                                        
                                        
                                        二〇二四年六月
"""
    
    def _get_english_cover_content(self):
        """生成英文封面内容"""
        return """
University Code: 10200                    Student ID: 【Student ID】
                                         Security Level: Public

Master's Thesis of Northeast Normal University

【English Thesis Title】

                                        
                                        Candidate: 【Name】
                                        Supervisor: 【Supervisor Name】, 【Title】
                                        Subject, Major: 【Major】
                                        Research Direction: 【Research Direction】
                                        
                                        
                                        
                                        June 2024
"""

    def _add_cover_page(self):
        """添加专门格式化的封面页"""
        if self.template_name != 'nenu_thesis':
            return
        
        # 添加分页符
        self.doc.add_page_break()
        
        # 学校代码和学号行
        info_para = self.doc.add_paragraph("学校代码：10200")
        info_para.add_run("                           ")
        info_para.add_run("研究生学号：【学号】")
        info_para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        for run in info_para.runs:
            run.font.name = '宋体'
            run.font.size = Pt(10.5)  # 五号
        info_para.paragraph_format.space_after = Pt(0)
        
        # 密级行
        security_para = self.doc.add_paragraph()
        security_para.add_run("                                        ")
        security_para.add_run("密级：公开")
        security_para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        for run in security_para.runs:
            run.font.name = '宋体'
            run.font.size = Pt(10.5)  # 五号
        security_para.paragraph_format.space_after = Pt(36)  # 3行间距
        
        # 空行
        self.doc.add_paragraph("")
        self.doc.add_paragraph("")
        
        # 大学名称
        univ_para = self.doc.add_paragraph("东北师范大学硕士学位论文")
        univ_para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        for run in univ_para.runs:
            run.font.name = '微软雅黑'
            run.font.size = Pt(16)  # 三号
            run.font.bold = True
        univ_para.paragraph_format.space_after = Pt(48)  # 4行间距
        
        # 空行
        self.doc.add_paragraph("")
        self.doc.add_paragraph("")
        
        # 论文题目
        title_para = self.doc.add_paragraph("【中文论文题目】")
        title_para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        for run in title_para.runs:
            run.font.name = '微软雅黑'
            run.font.size = Pt(18)  # 二号
            run.font.bold = True
        title_para.paragraph_format.space_after = Pt(84)  # 7行间距
        
        # 多个空行
        for _ in range(6):
            self.doc.add_paragraph("")
        
        # 学位申请人信息
        info_lines = [
            "学    位    申    请    人：【姓名】",
            "指    导    教    师：【导师姓名】  【职称】",
            "学    科、专    业：【专业名称】",
            "研    究    方    向：【研究方向】"
        ]
        
        for line in info_lines:
            info_para = self.doc.add_paragraph(line)
            info_para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            for run in info_para.runs:
                run.font.name = '宋体'
                run.font.size = Pt(12)  # 小四号
            info_para.paragraph_format.space_after = Pt(24)  # 2行间距
        
        # 多个空行
        for _ in range(8):
            self.doc.add_paragraph("")
        
        # 日期
        date_para = self.doc.add_paragraph("二〇二四年六月")
        date_para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        for run in date_para.runs:
            run.font.name = '宋体'
            run.font.size = Pt(14)  # 四号
        
        # 分页符
        self.doc.add_page_break()

    def _add_template_placeholders(self, missing_components):
        """为缺失的组件添加模板占位符"""
        if self.template_name != 'nenu_thesis':
            return
        
        # 添加封面页
        if 'cover_page' in missing_components:
            self._add_cover_page()
        
        # 定义组件占位符文本
        placeholders = {
            'cover_page': {
                'title': '东北师范大学硕士学位论文',
                'content': self._get_cover_content()
            },
            'english_cover': {
                'title': 'NENU Master\'s Thesis',
                'content': self._get_english_cover_content()
            },
            'declaration': {
                'title': '东北师范大学学位论文独创性声明',
                'content': '本人郑重声明：所呈交的学位论文是本人在导师指导下进行的研究工作及取得的研究成果。本论文除了文中特别加以标注和致谢的地方外，不包含其他人已经发表或撰写过的研究成果，也不包含为获得东北师范大学或其他教育机构的学位或证书而使用过的材料。与我一同工作的同志对本研究所做的任何贡献均已在论文中作了明确的说明并表示谢意。\n\n\n\n作者签名：【签名】\n\n日　　期：【日期】'
            },
            'authorization': {
                'title': '东北师范大学学位论文使用授权书',
                'content': '本人完全了解东北师范大学有关保留、使用学位论文的规定，同意学校保留并向国家有关部门或机构送交论文的复印件和电子版，允许论文被查阅和借阅。本人授权东北师范大学可以将本学位论文的全部或部分内容编入有关数据库进行检索，可以采用影印、缩印或扫描等复制手段保存和汇编本学位论文。\n\n\n\n作者签名：【签名】\n\n导师签名：【签名】\n\n日　　期：【日期】'
            },
            'abstract_cn': {
                'title': '摘　　要',  # 中间空2个汉字
                'content': '【此处为中文摘要内容，请根据实际情况填写】\n\n本论文...'
            },
            'abstract_en': {
                'title': 'Abstract',
                'content': '【Here is the English abstract content, please fill in according to the actual situation】\n\nThis thesis...'
            },
            'keywords_cn': {
                'content': '关键词：【关键词1；关键词2；关键词3】'
            },
            'keywords_en': {
                'content': 'Key words: 【keyword1; keyword2; keyword3】'
            },
            'toc': {
                'title': '目　　录',  # 中间空2个汉字
                'content': '【目录将在最终版本中自动生成】'
            },
            'symbols': {
                'title': '符号说明',
                'content': '【如有特殊符号，请在此处说明】'
            },
            'figures_list': {
                'title': '插图目录',
                'content': '【插图目录将在最终版本中自动生成】'
            },
            'tables_list': {
                'title': '附表目录',
                'content': '【附表目录将在最终版本中自动生成】'
            },
            'introduction': {
                'title': '第一章 引言',
                'content': '【请在此处写入引言内容】'
            },
            'literature_review': {
                'title': '第二章 文献综述',
                'content': '【请在此处写入文献综述内容】'
            },
            'methodology': {
                'title': '第三章 研究方法',
                'content': '【请在此处写入研究方法内容】'
            },
            'results': {
                'title': '第四章 研究结果',
                'content': '【请在此处写入研究结果内容】'
            },
            'discussion': {
                'title': '第五章 讨论',
                'content': '【请在此处写入讨论内容】'
            },
            'conclusion': {
                'title': '第六章 结论',
                'content': '【请在此处写入结论内容】'
            },
            'references': {
                'title': '参考文献',
                'content': '【请在此处添加参考文献】\n[1] 作者. 题目[M]. 出版地: 出版社, 年份.'
            },
            'appendix': {
                'title': '附录',
                'content': '【如有附录内容，请在此处添加】'
            },
            'acknowledgments': {
                'title': '致谢',
                'content': '【请在此处写入致谢内容】'
            }
        }
        
        # 为每个缺失的组件添加占位符
        for component in missing_components:
            if component in placeholders:
                placeholder = placeholders[component]
                
                # 添加标题（如果有）
                if 'title' in placeholder:
                    if component in ['declaration', 'authorization']:
                        # 独创性声明和授权书：三号黑体，居中
                        title_para = self.doc.add_paragraph(placeholder['title'])
                        title_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        for run in title_para.runs:
                            run.font.name = '黑体'
                            run.font.size = Pt(16)  # 三号
                            run.bold = True
                        title_para.paragraph_format.space_before = Pt(48)
                        title_para.paragraph_format.space_after = Pt(24)
                    elif component in ['abstract_cn', 'abstract_en']:
                        try:
                            style_name = 'Abstract Title CN' if component == 'abstract_cn' else 'Abstract Title EN'
                            title_para = self.doc.add_paragraph(placeholder['title'], style=style_name)
                        except:
                            title_para = self.doc.add_paragraph(placeholder['title'])
                            title_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    elif component == 'references':
                        try:
                            title_para = self.doc.add_paragraph(placeholder['title'], style='Reference Title')
                        except:
                            title_para = self.doc.add_paragraph(placeholder['title'])
                            title_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    elif component == 'toc':
                        try:
                            title_para = self.doc.add_paragraph(placeholder['title'], style='TOC Title')
                        except:
                            title_para = self.doc.add_paragraph(placeholder['title'])
                            title_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    elif component in ['introduction', 'literature_review', 'methodology', 'results', 'discussion', 'conclusion']:
                        self.doc.add_heading(placeholder['title'], level=1)
                    else:
                        # 其他组件使用一般标题格式
                        title_para = self.doc.add_paragraph(placeholder['title'])
                        title_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        for run in title_para.runs:
                            run.font.name = '黑体'
                            run.font.size = Pt(16)  # 三号
                            run.bold = True
                        title_para.paragraph_format.space_before = Pt(24)
                        title_para.paragraph_format.space_after = Pt(18)
                
                # 添加内容
                if 'content' in placeholder:
                    if component in ['declaration', 'authorization']:
                        # 独创性声明和授权书：宋体小四号，两端对齐，首行缩进2字符
                        content_para = self.doc.add_paragraph(placeholder['content'])
                        for run in content_para.runs:
                            run.font.name = '宋体'
                            run.font.size = Pt(12)  # 小四号
                        content_para.paragraph_format.first_line_indent = Cm(0.74)  # 首行缩进2字符
                        content_para.paragraph_format.line_spacing = 1.5
                        content_para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                    elif component in ['abstract_cn', 'abstract_en']:
                        try:
                            style_name = 'Abstract Body CN' if component == 'abstract_cn' else 'Abstract Body EN'
                            content_para = self.doc.add_paragraph(placeholder['content'], style=style_name)
                        except:
                            content_para = self.doc.add_paragraph(placeholder['content'])
                    elif component in ['keywords_cn', 'keywords_en']:
                        try:
                            style_name = 'Keywords CN' if component == 'keywords_cn' else 'Keywords EN'
                            content_para = self.doc.add_paragraph(placeholder['content'], style=style_name)
                        except:
                            content_para = self.doc.add_paragraph(placeholder['content'])
                    elif component == 'references':
                        try:
                            content_para = self.doc.add_paragraph(placeholder['content'], style='Reference Content')
                        except:
                            content_para = self.doc.add_paragraph(placeholder['content'])
                    elif component == 'appendix':
                        try:
                            content_para = self.doc.add_paragraph(placeholder['content'], style='Appendix Content')
                        except:
                            content_para = self.doc.add_paragraph(placeholder['content'])
                    else:
                        # 其他内容按正文格式
                        content_para = self.doc.add_paragraph(placeholder['content'])
                
                # 添加分页符（章节之间）
                if component in ['declaration', 'authorization', 'abstract_cn', 'abstract_en', 'toc', 'introduction', 'literature_review', 'methodology', 'results', 'discussion', 'conclusion', 'references', 'appendix', 'acknowledgments']:
                    self.doc.add_page_break()
    
    def _setup_page_numbering(self):
        """设置页码编排规则"""
        if self.template_name != 'nenu_thesis':
            return
            
        # 获取所有节
        sections = self.doc.sections
        
        for section in sections:
            # 设置页眉：东北师范大学硕士学位论文
            header = section.header
            if header.paragraphs:
                header_para = header.paragraphs[0]
                header_para.text = "东北师范大学硕士学位论文"
                header_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                
                # 设置页眉字体：小四号黑体
                for run in header_para.runs:
                    run.font.name = '黑体'
                    run.font.size = Pt(12)  # 小四号
            
            # 设置页脚页码：居中，五号Times New Roman
            footer = section.footer
            if footer.paragraphs:
                footer_para = footer.paragraphs[0]
                footer_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                
                # 清空原有内容
                footer_para.clear()
                
                # 添加页码
                run = footer_para.add_run()
                run.font.name = 'Times New Roman'
                run.font.size = Pt(10.5)  # 五号
                
                # 添加页码字段
                from docx.oxml import OxmlElement
                from docx.oxml.ns import qn
                
                fldChar1 = OxmlElement('w:fldChar')
                fldChar1.set(qn('w:fldCharType'), 'begin')
                run._element.append(fldChar1)
                
                instrText = OxmlElement('w:instrText')
                instrText.text = 'PAGE'
                run._element.append(instrText)
                
                fldChar2 = OxmlElement('w:fldChar')
                fldChar2.set(qn('w:fldCharType'), 'end')
                run._element.append(fldChar2)
    
    def batch_convert(self, input_dir, output_dir, use_pandoc=True):
        """批量转换目录中的所有Markdown文件"""
        input_path = Path(input_dir)
        output_path = Path(output_dir)
        
        # 创建输出目录
        output_path.mkdir(parents=True, exist_ok=True)
        
        # 查找所有Markdown文件
        md_files = list(input_path.glob('**/*.md'))
        
        if not md_files:
            print("未找到任何Markdown文件")
            return
        
        print(f"找到 {len(md_files)} 个Markdown文件")
        print(f"使用模板: {self.template.name}")
        
        success_count = 0
        for md_file in md_files:
            # 计算相对路径并创建对应的输出路径
            relative_path = md_file.relative_to(input_path)
            output_file = output_path / relative_path.with_suffix('.docx')
            
            # 创建输出文件的父目录
            output_file.parent.mkdir(parents=True, exist_ok=True)
            
            # 转换文件
            if use_pandoc:
                if self.convert_with_pandoc(str(md_file), str(output_file)):
                    success_count += 1
            else:
                if self.convert_with_python_docx(str(md_file), str(output_file)):
                    success_count += 1
        
        print(f"\n转换完成: 成功 {success_count}/{len(md_files)} 个文件")


def main():
    parser = argparse.ArgumentParser(description='Markdown转Word文档转换器')
    parser.add_argument('input', nargs='?', help='输入的Markdown文件或目录')
    parser.add_argument('-o', '--output', help='输出的Word文件或目录', default=None)
    parser.add_argument('--batch', action='store_true', help='批量转换模式')
    parser.add_argument('--method', choices=['pandoc', 'python-docx'], 
                       default='pandoc', help='转换方法（默认使用pandoc）')
    parser.add_argument('--template', choices=list(list_templates().keys()),
                       default='default', help='选择文档模板')
    parser.add_argument('--list-templates', action='store_true', 
                       help='列出所有可用模板')
    
    args = parser.parse_args()
    
    # 列出模板
    if args.list_templates:
        print("可用模板:")
        for name, description in list_templates().items():
            print(f"  {name}: {description}")
        return
    
    # 如果不是列出模板，则input参数是必需的
    if not args.input:
        parser.error("input参数是必需的（除非使用--list-templates）")
    
    converter = MarkdownToWordConverter(template_name=args.template)
    
    # 检查是否安装了pandoc
    if args.method == 'pandoc':
        try:
            pypandoc.get_pandoc_version()
        except Exception:
            print("警告: 未安装pandoc，将使用python-docx方法")
            print("建议安装pandoc以获得更好的转换效果:")
            print("  macOS: brew install pandoc")
            print("  Ubuntu: sudo apt-get install pandoc")
            print("  Windows: 从 https://pandoc.org/installing.html 下载安装")
            args.method = 'python-docx'
    
    use_pandoc = args.method == 'pandoc'
    
    if args.batch or os.path.isdir(args.input):
        # 批量转换模式
        output_dir = args.output or 'word_output'
        converter.batch_convert(args.input, output_dir, use_pandoc)
    else:
        # 单文件转换模式
        if not args.input.endswith('.md'):
            print("错误: 输入文件必须是.md文件")
            return
        
        output_file = args.output or args.input.replace('.md', '.docx')
        
        print(f"使用模板: {converter.template.name}")
        if use_pandoc:
            converter.convert_with_pandoc(args.input, output_file)
        else:
            converter.convert_with_python_docx(args.input, output_file)


if __name__ == '__main__':
    main()