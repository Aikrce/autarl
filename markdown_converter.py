#!/usr/bin/env python3
"""
Markdown to Word Converter - 完整版
支持表格、Mermaid图表、智能清理等功能
"""

import re
import os
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# 导入核心组件
from core.enhanced_table_converter import AdvancedTableConverter
from core.mermaid_converter import MermaidConverter


class MarkdownToWordConverter:
    """完整的Markdown到Word转换器"""
    
    def __init__(self, enable_mermaid=True, mermaid_method='web', auto_clean=True):
        """
        初始化转换器
        
        Args:
            enable_mermaid: 是否启用Mermaid图表支持
            mermaid_method: Mermaid转换方法 ('api', 'web', 'cli')
            auto_clean: 是否自动清理重复标题
        """
        self.enable_mermaid = enable_mermaid
        self.auto_clean = auto_clean
        self.table_converter = AdvancedTableConverter()
        
        if enable_mermaid:
            self.mermaid_converter = MermaidConverter(method=mermaid_method)
    
    def convert(self, input_file, output_file=None):
        """
        转换Markdown文件到Word
        
        Args:
            input_file: 输入的Markdown文件
            output_file: 输出的Word文件（可选）
        """
        if output_file is None:
            output_file = input_file.replace('.md', '.docx')
        
        # 读取文件
        with open(input_file, 'r', encoding='utf-8') as f:
            content = f.read()
        
        # 自动清理
        if self.auto_clean:
            content = self._clean_content(content)
        
        # 创建Word文档
        doc = Document()
        
        # 处理内容
        self._process_content(content, doc)
        
        # 保存文档
        doc.save(output_file)
        print(f"✅ 转换成功: {output_file}")
        
        return output_file
    
    def _clean_content(self, content):
        """清理重复的标题和格式"""
        lines = content.split('\n')
        cleaned_lines = []
        
        for i, line in enumerate(lines):
            stripped = line.strip()
            
            # 保留空行、分隔线、代码块
            if not stripped or stripped in ['---', '***', '___'] or stripped.startswith('```'):
                cleaned_lines.append(line)
                continue
            
            # 删除只有Markdown标记的标题行
            if stripped.startswith('#') and not re.search(r'第[一二三四五六七八九十]+[章节]', stripped):
                # 检查是否只有标记没有内容
                title_text = stripped.lstrip('#').strip()
                if not title_text:
                    continue
                
                # 检查下一行是否包含相同内容
                if i + 1 < len(lines):
                    next_line = lines[i + 1].strip()
                    if title_text in next_line:
                        continue
            
            # 处理双重编号
            if re.match(r'^（[^）]+）（[^）]+）', stripped):
                line = re.sub(r'^（[^）]+）', '', stripped)
            
            cleaned_lines.append(line)
        
        return '\n'.join(cleaned_lines)
    
    def _process_content(self, content, doc):
        """处理Markdown内容"""
        lines = content.split('\n')
        i = 0
        
        while i < len(lines):
            line = lines[i]
            
            # Mermaid代码块
            if self.enable_mermaid and line.strip() == '```mermaid':
                mermaid_code = self._extract_code_block(lines, i + 1)
                if mermaid_code:
                    self.mermaid_converter.add_mermaid_to_document(doc, mermaid_code)
                    i = self._skip_code_block(lines, i + 1)
                    continue
            
            # 表格
            if '|' in line and not line.strip().startswith('```'):
                table_data = self._extract_table(lines, i)
                if table_data:
                    self.table_converter.add_table_to_document(doc, {'data': table_data})
                    i = self._skip_table(lines, i)
                    continue
            
            # 其他元素
            self._process_line(doc, line)
            i += 1
    
    def _process_line(self, doc, line):
        """处理单行内容"""
        stripped = line.strip()
        
        # 分隔线
        if stripped in ['---', '***', '___']:
            doc.add_page_break()
            return
        
        # 空行
        if not stripped:
            doc.add_paragraph()
            return
        
        # 标题
        if stripped.startswith('#'):
            level = len(stripped) - len(stripped.lstrip('#'))
            title = stripped.lstrip('#').strip()
            if title:
                doc.add_heading(title, level=min(level, 6))
            return
        
        # 列表
        if stripped.startswith(('- ', '* ', '+ ')):
            doc.add_paragraph(stripped[2:], style='List Bullet')
            return
        
        # 普通段落
        doc.add_paragraph(stripped)
    
    def _extract_code_block(self, lines, start):
        """提取代码块内容"""
        code_lines = []
        i = start
        while i < len(lines) and not lines[i].strip().startswith('```'):
            code_lines.append(lines[i])
            i += 1
        return '\n'.join(code_lines) if code_lines else None
    
    def _skip_code_block(self, lines, start):
        """跳过代码块"""
        i = start
        while i < len(lines) and not lines[i].strip().startswith('```'):
            i += 1
        return i + 1
    
    def _extract_table(self, lines, start):
        """提取表格数据"""
        table_lines = []
        i = start
        
        while i < len(lines) and '|' in lines[i]:
            line = lines[i].strip()
            if not re.match(r'^\|[\s\-\|:]+\|?$', line):
                cells = [cell.strip() for cell in line.split('|')]
                cells = [c for c in cells if c]
                if cells:
                    table_lines.append(cells)
            i += 1
        
        return table_lines if table_lines else None
    
    def _skip_table(self, lines, start):
        """跳过表格"""
        i = start
        while i < len(lines) and '|' in lines[i]:
            i += 1
        return i


def main():
    """主函数"""
    import argparse
    
    parser = argparse.ArgumentParser(description='Markdown to Word Converter')
    parser.add_argument('input', help='Input markdown file')
    parser.add_argument('-o', '--output', help='Output word file')
    parser.add_argument('--no-mermaid', action='store_true', help='Disable Mermaid support')
    parser.add_argument('--no-clean', action='store_true', help='Disable auto cleaning')
    
    args = parser.parse_args()
    
    converter = MarkdownToWordConverter(
        enable_mermaid=not args.no_mermaid,
        auto_clean=not args.no_clean
    )
    
    converter.convert(args.input, args.output)


if __name__ == '__main__':
    main()
