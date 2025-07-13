#!/usr/bin/env python3
"""
Standard Academic Paper Template Generator
标准学术论文模板生成器
"""

import os
import json
import logging
from typing import Dict, List, Optional, Any
from pathlib import Path
from datetime import datetime

from docx import Document
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.shared import OxmlElement, qn

from word_template_analyzer import TemplateLibrary, analyze_word_template

logger = logging.getLogger(__name__)


class StandardAcademicTemplateGenerator:
    """标准学术论文模板生成器"""
    
    def __init__(self, template_library_path: str = "template_library"):
        self.template_library = TemplateLibrary(template_library_path)
        self.library_path = template_library_path
        
    def create_standard_academic_template(self, 
                                        template_name: str = "标准学术论文模板") -> str:
        """
        创建标准学术论文模板
        
        Args:
            template_name: 模板名称
            
        Returns:
            str: 新模板ID
        """
        try:
            # 创建新的Word文档
            doc = Document()
            
            # 设置页面布局
            self._setup_page_layout(doc)
            
            # 创建标准学术样式
            self._create_academic_styles(doc)
            
            # 添加示例内容结构
            self._add_template_structure(doc)
            
            # 生成模板ID
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            template_id = f"standard_academic_{timestamp}"
            
            # 创建模板目录
            template_dir = os.path.join(self.library_path, template_id)
            os.makedirs(template_dir, exist_ok=True)
            
            # 保存模板文件
            template_path = os.path.join(template_dir, f"{template_id}.docx")
            doc.save(template_path)
            
            # 分析模板
            template_info = analyze_word_template(template_path)
            
            # 保存配置
            config_path = os.path.join(template_dir, "template_config.json")
            with open(config_path, 'w', encoding='utf-8') as f:
                json.dump(template_info.to_dict(), f, ensure_ascii=False, indent=2)
            
            # 创建内容结构配置
            structure_config = self._create_structure_config()
            structure_path = os.path.join(template_dir, "content_structure.json")
            with open(structure_path, 'w', encoding='utf-8') as f:
                json.dump(structure_config, f, ensure_ascii=False, indent=2)
            
            # 更新模板库索引
            self._update_template_index(template_id, template_name, template_path,
                                      config_path, structure_path, template_info)
            
            logger.info(f"标准学术论文模板创建成功: {template_id}")
            return template_id
            
        except Exception as e:
            logger.error(f"创建标准学术论文模板失败: {e}")
            raise
    
    def _setup_page_layout(self, doc: Document):
        """设置页面布局"""
        # 设置页面尺寸为A4
        section = doc.sections[0]
        section.page_width = Cm(21)
        section.page_height = Cm(29.7)
        
        # 设置页边距
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.0)
        
        # 设置页眉页脚距离
        section.header_distance = Cm(1.5)
        section.footer_distance = Cm(1.25)
        
        logger.info("页面布局设置完成")
    
    def _create_academic_styles(self, doc: Document):
        """创建标准学术样式"""
        styles = doc.styles
        
        # 1. 标题样式
        self._create_title_style(styles)
        
        # 2. 正文样式
        self._create_normal_style(styles)
        
        # 3. 各级标题样式
        self._create_heading_styles(styles)
        
        # 4. 参考文献样式
        self._create_reference_style(styles)
        
        # 5. 脚注样式
        self._create_footnote_style(styles)
        
        # 6. 图表标题样式
        self._create_caption_styles(styles)
        
        # 7. 摘要样式
        self._create_abstract_styles(styles)
        
        logger.info("学术样式创建完成")
    
    def _create_title_style(self, styles):
        """创建标题样式"""
        try:
            title_style = styles.add_style('Paper Title', WD_STYLE_TYPE.PARAGRAPH)
            title_style.font.name = 'Times New Roman'
            title_style.font.size = Pt(16)
            title_style.font.bold = True
            title_style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            title_style.paragraph_format.space_before = Pt(0)
            title_style.paragraph_format.space_after = Pt(18)
            title_style.paragraph_format.line_spacing = 1.15
        except Exception as e:
            logger.warning(f"创建标题样式失败: {e}")
    
    def _create_normal_style(self, styles):
        """创建正文样式"""
        try:
            # 修改默认正文样式
            normal_style = styles['Normal']
            normal_style.font.name = 'Times New Roman'
            normal_style.font.size = Pt(12)
            normal_style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
            normal_style.paragraph_format.line_spacing = 1.5
            normal_style.paragraph_format.first_line_indent = Cm(0.75)
            normal_style.paragraph_format.space_before = Pt(0)
            normal_style.paragraph_format.space_after = Pt(0)
        except Exception as e:
            logger.warning(f"创建正文样式失败: {e}")
    
    def _create_heading_styles(self, styles):
        """创建各级标题样式"""
        heading_configs = [
            {
                'name': 'Heading 1',
                'font_size': 16,
                'bold': True,
                'alignment': WD_PARAGRAPH_ALIGNMENT.CENTER,
                'space_before': 18,
                'space_after': 12,
                'numbering': True
            },
            {
                'name': 'Heading 2', 
                'font_size': 14,
                'bold': True,
                'alignment': WD_PARAGRAPH_ALIGNMENT.LEFT,
                'space_before': 12,
                'space_after': 6,
                'first_line_indent': 0.75
            },
            {
                'name': 'Heading 3',
                'font_size': 12,
                'bold': True,
                'alignment': WD_PARAGRAPH_ALIGNMENT.LEFT,
                'space_before': 6,
                'space_after': 3,
                'first_line_indent': 0.75
            },
            {
                'name': 'Heading 4',
                'font_size': 12,
                'bold': False,
                'alignment': WD_PARAGRAPH_ALIGNMENT.LEFT,
                'space_before': 6,
                'space_after': 3,
                'first_line_indent': 0.75
            }
        ]
        
        for config in heading_configs:
            try:
                if config['name'] in styles:
                    style = styles[config['name']]
                else:
                    style = styles.add_style(config['name'], WD_STYLE_TYPE.PARAGRAPH)
                
                style.font.name = 'Times New Roman'
                style.font.size = Pt(config['font_size'])
                style.font.bold = config['bold']
                style.paragraph_format.alignment = config['alignment']
                style.paragraph_format.space_before = Pt(config['space_before'])
                style.paragraph_format.space_after = Pt(config['space_after'])
                style.paragraph_format.line_spacing = 1.5
                
                if 'first_line_indent' in config:
                    style.paragraph_format.first_line_indent = Cm(config['first_line_indent'])
                    
            except Exception as e:
                logger.warning(f"创建标题样式 {config['name']} 失败: {e}")
    
    def _create_reference_style(self, styles):
        """创建参考文献样式"""
        try:
            ref_style = styles.add_style('References', WD_STYLE_TYPE.PARAGRAPH)
            ref_style.font.name = 'Times New Roman'
            ref_style.font.size = Pt(10.5)
            ref_style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
            ref_style.paragraph_format.line_spacing = 1.5
            ref_style.paragraph_format.hanging_indent = Cm(0.75)
            ref_style.paragraph_format.space_before = Pt(0)
            ref_style.paragraph_format.space_after = Pt(3)
        except Exception as e:
            logger.warning(f"创建参考文献样式失败: {e}")
    
    def _create_footnote_style(self, styles):
        """创建脚注样式"""
        try:
            footnote_style = styles.add_style('Academic Footnote', WD_STYLE_TYPE.PARAGRAPH)
            footnote_style.font.name = 'Times New Roman'
            footnote_style.font.size = Pt(9)
            footnote_style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
            footnote_style.paragraph_format.line_spacing = 1.0
            footnote_style.paragraph_format.hanging_indent = Cm(0.35)
            footnote_style.paragraph_format.space_before = Pt(0)
            footnote_style.paragraph_format.space_after = Pt(0)
        except Exception as e:
            logger.warning(f"创建脚注样式失败: {e}")
    
    def _create_caption_styles(self, styles):
        """创建图表标题样式"""
        try:
            # 表格标题
            table_caption = styles.add_style('Table Caption', WD_STYLE_TYPE.PARAGRAPH)
            table_caption.font.name = 'Times New Roman'
            table_caption.font.size = Pt(10.5)
            table_caption.font.bold = True
            table_caption.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            table_caption.paragraph_format.space_before = Pt(6)
            table_caption.paragraph_format.space_after = Pt(6)
            table_caption.paragraph_format.line_spacing = 1.15
            
            # 图片标题
            figure_caption = styles.add_style('Figure Caption', WD_STYLE_TYPE.PARAGRAPH)
            figure_caption.font.name = 'Times New Roman'
            figure_caption.font.size = Pt(10.5)
            figure_caption.font.bold = True
            figure_caption.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            figure_caption.paragraph_format.space_before = Pt(6)
            figure_caption.paragraph_format.space_after = Pt(12)
            figure_caption.paragraph_format.line_spacing = 1.15
            
        except Exception as e:
            logger.warning(f"创建图表标题样式失败: {e}")
    
    def _create_abstract_styles(self, styles):
        """创建摘要样式"""
        try:
            # 摘要标题
            abstract_title = styles.add_style('Abstract Title', WD_STYLE_TYPE.PARAGRAPH)
            abstract_title.font.name = 'Times New Roman'
            abstract_title.font.size = Pt(14)
            abstract_title.font.bold = True
            abstract_title.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            abstract_title.paragraph_format.space_before = Pt(18)
            abstract_title.paragraph_format.space_after = Pt(12)
            abstract_title.paragraph_format.line_spacing = 1.15
            
            # 摘要正文
            abstract_body = styles.add_style('Abstract Body', WD_STYLE_TYPE.PARAGRAPH)
            abstract_body.font.name = 'Times New Roman'
            abstract_body.font.size = Pt(11)
            abstract_body.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
            abstract_body.paragraph_format.line_spacing = 1.15
            abstract_body.paragraph_format.space_before = Pt(0)
            abstract_body.paragraph_format.space_after = Pt(0)
            abstract_body.paragraph_format.left_indent = Cm(0.5)
            abstract_body.paragraph_format.right_indent = Cm(0.5)
            
            # 关键词
            keywords = styles.add_style('Keywords', WD_STYLE_TYPE.PARAGRAPH)
            keywords.font.name = 'Times New Roman'
            keywords.font.size = Pt(11)
            keywords.font.bold = True
            keywords.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
            keywords.paragraph_format.line_spacing = 1.15
            keywords.paragraph_format.space_before = Pt(6)
            keywords.paragraph_format.space_after = Pt(12)
            keywords.paragraph_format.left_indent = Cm(0.5)
            keywords.paragraph_format.right_indent = Cm(0.5)
            
        except Exception as e:
            logger.warning(f"创建摘要样式失败: {e}")
    
    def _add_template_structure(self, doc: Document):
        """添加模板结构示例"""
        # 标题
        title_para = doc.add_paragraph("学术论文标题", style='Paper Title')
        
        # 作者信息
        author_para = doc.add_paragraph("作者姓名", style='Normal')
        author_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        author_para.paragraph_format.space_after = Pt(6)
        
        affiliation_para = doc.add_paragraph("作者单位", style='Normal')
        affiliation_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        affiliation_para.paragraph_format.space_after = Pt(18)
        
        # 中文摘要
        doc.add_paragraph("摘要", style='Abstract Title')
        doc.add_paragraph("这里是中文摘要内容。摘要应简明扼要地概述研究目的、方法、结果和结论。", style='Abstract Body')
        doc.add_paragraph("关键词：学术论文；模板；格式规范", style='Keywords')
        
        # 英文摘要
        doc.add_paragraph("Abstract", style='Abstract Title')
        doc.add_paragraph("This is the English abstract content. The abstract should briefly summarize the research purpose, methods, results, and conclusions.", style='Abstract Body')
        doc.add_paragraph("Keywords: academic paper; template; formatting standards", style='Keywords')
        
        # 目录占位符
        doc.add_page_break()
        doc.add_paragraph("目录", style='Heading 1')
        doc.add_paragraph("（此处将自动生成目录）", style='Normal')
        
        # 正文结构
        doc.add_page_break()
        doc.add_paragraph("1. 引言", style='Heading 1')
        doc.add_paragraph("这是引言部分的正文内容。引言应介绍研究背景、目的和意义。", style='Normal')
        
        doc.add_paragraph("1.1 研究背景", style='Heading 2')
        doc.add_paragraph("这是二级标题下的正文内容。", style='Normal')
        
        doc.add_paragraph("1.1.1 具体问题", style='Heading 3')
        doc.add_paragraph("这是三级标题下的正文内容。", style='Normal')
        
        doc.add_paragraph("2. 文献综述", style='Heading 1')
        doc.add_paragraph("这是文献综述部分的正文内容。", style='Normal')
        
        doc.add_paragraph("3. 研究方法", style='Heading 1')
        doc.add_paragraph("这是研究方法部分的正文内容。", style='Normal')
        
        doc.add_paragraph("4. 结果与分析", style='Heading 1')
        doc.add_paragraph("这是结果与分析部分的正文内容。", style='Normal')
        
        # 表格示例
        doc.add_paragraph("表1 研究结果汇总", style='Table Caption')
        table = doc.add_table(rows=3, cols=3)
        table.style = 'Table Grid'
        
        # 图片示例
        doc.add_paragraph("图1 研究框架图", style='Figure Caption')
        
        doc.add_paragraph("5. 讨论", style='Heading 1')
        doc.add_paragraph("这是讨论部分的正文内容。", style='Normal')
        
        doc.add_paragraph("6. 结论", style='Heading 1')
        doc.add_paragraph("这是结论部分的正文内容。", style='Normal')
        
        # 参考文献
        doc.add_page_break()
        doc.add_paragraph("参考文献", style='Heading 1')
        doc.add_paragraph("[1] 作者. 文献标题[J]. 期刊名称, 年份, 卷(期): 页码.", style='References')
        doc.add_paragraph("[2] 作者. 书籍标题[M]. 出版地: 出版社, 年份.", style='References')
        
        logger.info("模板结构添加完成")
    
    def _create_structure_config(self) -> Dict[str, Any]:
        """创建内容结构配置"""
        return {
            "document_structure": {
                "title_page": True,
                "abstract_cn": True,
                "abstract_en": True,
                "keywords": True,
                "table_of_contents": True,
                "main_content": True,
                "references": True,
                "appendix": False,
                "acknowledgments": False
            },
            "formatting_rules": {
                "page_format": "A4",
                "margins": "2.5cm(左右) 2.5cm(上) 2.0cm(下)",
                "line_spacing": "1.5倍行距",
                "font_family": "Times New Roman",
                "font_size": "12pt",
                "heading_numbering": "numeric",
                "reference_style": "academic_standard"
            },
            "style_features": {
                "title_centered": True,
                "first_line_indent": "0.75cm",
                "justified_text": True,
                "hanging_indent_references": True,
                "footnote_support": True,
                "caption_numbering": True
            },
            "template_info": {
                "created_at": datetime.now().isoformat(),
                "template_type": "standard_academic",
                "version": "1.0",
                "description": "标准学术论文模板，适用于各类学术论文写作"
            }
        }
    
    def _update_template_index(self, template_id: str, name: str, word_file: str,
                             config_file: str, structure_file: str, template_info) -> None:
        """更新模板库索引"""
        try:
            index_path = os.path.join(self.library_path, "template_index.json")
            
            # 读取现有索引
            if os.path.exists(index_path):
                with open(index_path, 'r', encoding='utf-8') as f:
                    index_data = json.load(f)
            else:
                index_data = {"templates": {}, "version": "1.0"}
            
            # 添加新模板
            index_data["templates"][template_id] = {
                "name": name,
                "description": "标准学术论文模板，适用于各类学术论文写作，包含完整的论文结构和格式规范",
                "tags": ["academic", "standard", "thesis", "paper", "universal"],
                "created_at": datetime.now().isoformat(),
                "word_file": word_file,
                "config_file": config_file,
                "structure_file": structure_file,
                "styles_count": len(template_info.styles) if hasattr(template_info, 'styles') else 0,
                "page_setup": {
                    "width": 21.0,
                    "height": 29.7,
                    "orientation": "portrait"
                },
                "is_standard": True,
                "template_type": "academic_standard",
                "priority": "standard"
            }
            
            # 保存索引
            with open(index_path, 'w', encoding='utf-8') as f:
                json.dump(index_data, f, ensure_ascii=False, indent=2)
                
        except Exception as e:
            logger.error(f"更新模板索引失败: {e}")
            raise


def create_standard_academic_template(template_library_path: str = "template_library") -> str:
    """创建标准学术论文模板的便捷函数"""
    generator = StandardAcademicTemplateGenerator(template_library_path)
    return generator.create_standard_academic_template()


if __name__ == "__main__":
    logging.basicConfig(level=logging.INFO)
    
    try:
        template_id = create_standard_academic_template()
        print(f"✅ 标准学术论文模板创建成功！模板ID: {template_id}")
    except Exception as e:
        print(f"❌ 创建标准学术论文模板失败: {e}")