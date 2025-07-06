#!/usr/bin/env python3

import os
import argparse
from pathlib import Path
import pypandoc
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
import markdown2
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from templates_config import get_template, list_templates


class MarkdownToWordConverter:
    def __init__(self, template_name='default'):
        self.doc = None
        self.template_name = template_name
        self.template = get_template(template_name)
        
    def convert_with_pandoc(self, input_file, output_file, output_format='docx'):
        """使用pandoc进行转换（推荐方式）"""
        try:
            extra_args = ['--standalone']
            if output_format == 'docx':
                extra_args.append('--toc')
            elif output_format == 'pdf':
                extra_args.extend(['--pdf-engine=xelatex', '--toc'])
            elif output_format == 'html':
                extra_args.extend(['--toc', '--css=style.css'])
            
            pypandoc.convert_file(
                input_file,
                output_format,
                outputfile=output_file,
                extra_args=extra_args
            )
            print(f"✓ 成功转换: {input_file} -> {output_file}")
            return True
        except Exception as e:
            print(f"✗ Pandoc转换失败: {e}")
            return False
    
    def convert_to_html(self, input_file, output_file):
        """转换为HTML格式"""
        try:
            with open(input_file, 'r', encoding='utf-8') as f:
                markdown_text = f.read()
            
            html = markdown2.markdown(
                markdown_text,
                extras=['tables', 'fenced-code-blocks', 'header-ids', 'toc']
            )
            
            html_template = f"""<!DOCTYPE html>
<html lang="zh-CN">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Markdown Document</title>
    <style>
        body {{
            font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', 'Microsoft YaHei', sans-serif;
            line-height: 1.6;
            color: #333;
            max-width: 800px;
            margin: 0 auto;
            padding: 20px;
        }}
        code {{
            background-color: #f4f4f4;
            padding: 2px 4px;
            border-radius: 3px;
        }}
        pre {{
            background-color: #f4f4f4;
            padding: 10px;
            border-radius: 5px;
            overflow-x: auto;
        }}
        blockquote {{
            border-left: 4px solid #ddd;
            padding-left: 16px;
            margin-left: 0;
        }}
        table {{
            border-collapse: collapse;
            width: 100%;
        }}
        th, td {{
            border: 1px solid #ddd;
            padding: 8px;
            text-align: left;
        }}
        th {{
            background-color: #f2f2f2;
        }}
    </style>
</head>
<body>
{html}
</body>
</html>"""
            
            with open(output_file, 'w', encoding='utf-8') as f:
                f.write(html_template)
            
            print(f"✓ 成功转换: {input_file} -> {output_file}")
            return True
            
        except Exception as e:
            print(f"✗ HTML转换失败: {e}")
            return False
    
    def convert_to_txt(self, input_file, output_file):
        """转换为纯文本格式"""
        try:
            with open(input_file, 'r', encoding='utf-8') as f:
                markdown_text = f.read()
            
            # 简单的Markdown到纯文本转换
            import re
            
            # 移除图片链接
            text = re.sub(r'!\[.*?\]\(.*?\)', '', markdown_text)
            # 移除链接，保留链接文本
            text = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', text)
            # 移除代码块标记
            text = re.sub(r'```[^`]*```', '', text, flags=re.DOTALL)
            # 移除行内代码标记
            text = re.sub(r'`([^`]+)`', r'\1', text)
            # 移除粗体和斜体标记
            text = re.sub(r'\*\*([^*]+)\*\*', r'\1', text)
            text = re.sub(r'__([^_]+)__', r'\1', text)
            text = re.sub(r'\*([^*]+)\*', r'\1', text)
            text = re.sub(r'_([^_]+)_', r'\1', text)
            # 移除标题标记
            text = re.sub(r'^#+\s*', '', text, flags=re.MULTILINE)
            # 移除列表标记
            text = re.sub(r'^[\s]*[-*+]\s*', '', text, flags=re.MULTILINE)
            text = re.sub(r'^\s*\d+\.\s*', '', text, flags=re.MULTILINE)
            
            with open(output_file, 'w', encoding='utf-8') as f:
                f.write(text)
            
            print(f"✓ 成功转换: {input_file} -> {output_file}")
            return True
            
        except Exception as e:
            print(f"✗ TXT转换失败: {e}")
            return False
    
    def convert_with_python_docx(self, input_file, output_file):
        """使用python-docx进行转换（备用方式）"""
        try:
            with open(input_file, 'r', encoding='utf-8') as f:
                markdown_text = f.read()
            
            # 使用markdown2解析Markdown
            html = markdown2.markdown(
                markdown_text,
                extras=['tables', 'fenced-code-blocks', 'header-ids']
            )
            
            # 创建Word文档
            self.doc = Document()
            self._setup_styles()
            
            # 解析并添加内容
            self._parse_markdown_content(markdown_text)
            
            # 保存文档
            self.doc.save(output_file)
            print(f"✓ 成功转换: {input_file} -> {output_file}")
            return True
            
        except Exception as e:
            print(f"✗ Python-docx转换失败: {e}")
            return False
    
    def _setup_styles(self):
        """设置文档样式"""
        # 应用选定的模板
        self.template.apply_to_document(self.doc)
    
    def _parse_markdown_content(self, content):
        """解析Markdown内容并添加到Word文档"""
        lines = content.split('\n')
        current_paragraph = None
        in_code_block = False
        code_content = []
        
        for line in lines:
            # 代码块处理
            if line.strip().startswith('```'):
                if in_code_block:
                    # 结束代码块
                    code_text = '\n'.join(code_content)
                    p = self.doc.add_paragraph()
                    run = p.add_run(code_text)
                    run.font.name = 'Consolas'
                    run.font.size = Pt(10)
                    run.font.color.rgb = RGBColor(0, 0, 0)
                    # 设置段落样式
                    p.paragraph_format.left_indent = Pt(18)
                    self._add_border(p)
                    code_content = []
                    in_code_block = False
                else:
                    # 开始代码块
                    in_code_block = True
                continue
            
            if in_code_block:
                code_content.append(line)
                continue
            
            # 标题处理
            if line.startswith('#'):
                level = len(line) - len(line.lstrip('#'))
                if level <= 6:
                    title_text = line.lstrip('#').strip()
                    self.doc.add_heading(title_text, level=level)
                    continue
            
            # 列表处理
            if line.strip().startswith(('- ', '* ', '+ ')):
                list_text = line.strip()[2:]
                p = self.doc.add_paragraph(list_text, style='List Bullet')
                continue
            
            if line.strip() and line.strip()[0].isdigit() and line.strip().find('.') == 1:
                list_text = line.strip()[3:]
                p = self.doc.add_paragraph(list_text, style='List Number')
                continue
            
            # 空行处理
            if not line.strip():
                if current_paragraph:
                    current_paragraph = None
                continue
            
            # 普通段落
            if current_paragraph is None:
                current_paragraph = self.doc.add_paragraph()
            
            # 处理行内格式
            self._process_inline_formatting(current_paragraph, line)
    
    def _process_inline_formatting(self, paragraph, text):
        """处理行内格式（粗体、斜体、代码等）"""
        import re
        
        # 简单的格式处理
        parts = re.split(r'(\*\*.*?\*\*|__.*?__|_.*?_|\*.*?\*|`.*?`)', text)
        
        for part in parts:
            if not part:
                continue
                
            # 粗体
            if part.startswith('**') and part.endswith('**'):
                run = paragraph.add_run(part[2:-2])
                run.bold = True
            elif part.startswith('__') and part.endswith('__'):
                run = paragraph.add_run(part[2:-2])
                run.bold = True
            # 斜体
            elif part.startswith('*') and part.endswith('*') and len(part) > 2:
                run = paragraph.add_run(part[1:-1])
                run.italic = True
            elif part.startswith('_') and part.endswith('_') and len(part) > 2:
                run = paragraph.add_run(part[1:-1])
                run.italic = True
            # 行内代码
            elif part.startswith('`') and part.endswith('`'):
                run = paragraph.add_run(part[1:-1])
                run.font.name = 'Consolas'
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(219, 48, 105)
            else:
                paragraph.add_run(part)
    
    def _add_border(self, paragraph):
        """为段落添加边框"""
        pPr = paragraph._p.get_or_add_pPr()
        borders = OxmlElement('w:pBdr')
        pPr.insert_element_before(borders,
            'w:shd', 'w:tabs', 'w:suppressAutoHyphens', 'w:kinsoku',
            'w:wordWrap', 'w:overflowPunct', 'w:topLinePunct', 'w:autoSpaceDE',
            'w:autoSpaceDN', 'w:bidi', 'w:adjustRightInd', 'w:snapToGrid',
            'w:spacing', 'w:ind', 'w:contextualSpacing', 'w:mirrorIndents',
            'w:suppressOverlap', 'w:jc', 'w:textDirection', 'w:textAlignment',
            'w:textboxTightWrap', 'w:outlineLvl', 'w:divId', 'w:cnfStyle',
            'w:rPr', 'w:sectPr', 'w:pPrChange'
        )
        
        for border_name in ['top', 'left', 'bottom', 'right']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), '4')
            border.set(qn('w:space'), '1')
            border.set(qn('w:color'), 'auto')
            borders.append(border)
    
    def batch_convert(self, input_dir, output_dir, use_pandoc=True):
        """批量转换目录中的所有Markdown文件"""
        input_path = Path(input_dir)
        output_path = Path(output_dir)
        
        # 创建输出目录
        output_path.mkdir(parents=True, exist_ok=True)
        
        # 查找所有Markdown文件
        md_files = list(input_path.glob('**/*.md'))
        
        if not md_files:
            print("未找到任何Markdown文件")
            return
        
        print(f"找到 {len(md_files)} 个Markdown文件")
        print(f"使用模板: {self.template.name}")
        
        success_count = 0
        for md_file in md_files:
            # 计算相对路径并创建对应的输出路径
            relative_path = md_file.relative_to(input_path)
            output_file = output_path / relative_path.with_suffix('.docx')
            
            # 创建输出文件的父目录
            output_file.parent.mkdir(parents=True, exist_ok=True)
            
            # 转换文件
            if use_pandoc:
                if self.convert_with_pandoc(str(md_file), str(output_file)):
                    success_count += 1
            else:
                if self.convert_with_python_docx(str(md_file), str(output_file)):
                    success_count += 1
        
        print(f"\n转换完成: 成功 {success_count}/{len(md_files)} 个文件")


def main():
    parser = argparse.ArgumentParser(description='Markdown转Word文档转换器')
    parser.add_argument('input', nargs='?', help='输入的Markdown文件或目录')
    parser.add_argument('-o', '--output', help='输出的Word文件或目录', default=None)
    parser.add_argument('--batch', action='store_true', help='批量转换模式')
    parser.add_argument('--method', choices=['pandoc', 'python-docx'], 
                       default='pandoc', help='转换方法（默认使用pandoc）')
    parser.add_argument('--template', choices=list(list_templates().keys()),
                       default='default', help='选择文档模板')
    parser.add_argument('--list-templates', action='store_true', 
                       help='列出所有可用模板')
    
    args = parser.parse_args()
    
    # 列出模板
    if args.list_templates:
        print("可用模板:")
        for name, description in list_templates().items():
            print(f"  {name}: {description}")
        return
    
    # 如果不是列出模板，则input参数是必需的
    if not args.input:
        parser.error("input参数是必需的（除非使用--list-templates）")
    
    converter = MarkdownToWordConverter(template_name=args.template)
    
    # 检查是否安装了pandoc
    if args.method == 'pandoc':
        try:
            pypandoc.get_pandoc_version()
        except Exception:
            print("警告: 未安装pandoc，将使用python-docx方法")
            print("建议安装pandoc以获得更好的转换效果:")
            print("  macOS: brew install pandoc")
            print("  Ubuntu: sudo apt-get install pandoc")
            print("  Windows: 从 https://pandoc.org/installing.html 下载安装")
            args.method = 'python-docx'
    
    use_pandoc = args.method == 'pandoc'
    
    if args.batch or os.path.isdir(args.input):
        # 批量转换模式
        output_dir = args.output or 'word_output'
        converter.batch_convert(args.input, output_dir, use_pandoc)
    else:
        # 单文件转换模式
        if not args.input.endswith('.md'):
            print("错误: 输入文件必须是.md文件")
            return
        
        output_file = args.output or args.input.replace('.md', '.docx')
        
        print(f"使用模板: {converter.template.name}")
        if use_pandoc:
            converter.convert_with_pandoc(args.input, output_file)
        else:
            converter.convert_with_python_docx(args.input, output_file)


if __name__ == '__main__':
    main()