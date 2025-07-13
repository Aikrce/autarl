#!/usr/bin/env python3
"""
Word Template Analyzer
Word模板分析器 - 解析上传的Word文档，提取完整的样式和格式信息
"""

import os
import logging
from typing import Dict, List, Optional, Any, Tuple
from dataclasses import dataclass, field
from pathlib import Path
import json
import zipfile
import xml.etree.ElementTree as ET
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
import re

logger = logging.getLogger(__name__)


@dataclass
class WordStyleInfo:
    """Word样式信息"""
    style_id: str
    name: str
    style_type: str
    
    # 字体信息
    font_name: Optional[str] = None
    font_size: Optional[float] = None
    font_bold: bool = False
    font_italic: bool = False
    font_underline: bool = False
    font_color: Optional[str] = None
    
    # 段落信息
    alignment: Optional[str] = None
    space_before: Optional[float] = None
    space_after: Optional[float] = None
    line_spacing: Optional[float] = None
    first_line_indent: Optional[float] = None
    left_indent: Optional[float] = None
    right_indent: Optional[float] = None
    hanging_indent: Optional[float] = None
    
    # 编号和列表
    numbering_id: Optional[str] = None
    numbering_level: Optional[int] = None
    
    # 边框和底纹
    borders: Dict[str, Any] = field(default_factory=dict)
    shading: Dict[str, Any] = field(default_factory=dict)
    
    # 基于的样式
    based_on: Optional[str] = None
    next_style: Optional[str] = None
    
    def to_dict(self) -> Dict[str, Any]:
        """转换为字典格式"""
        return {
            'style_id': self.style_id,
            'name': self.name,
            'style_type': self.style_type,
            'font': {
                'name': self.font_name,
                'size': self.font_size,
                'bold': self.font_bold,
                'italic': self.font_italic,
                'underline': self.font_underline,
                'color': self.font_color
            },
            'paragraph': {
                'alignment': self.alignment,
                'space_before': self.space_before,
                'space_after': self.space_after,
                'line_spacing': self.line_spacing,
                'first_line_indent': self.first_line_indent,
                'left_indent': self.left_indent,
                'right_indent': self.right_indent,
                'hanging_indent': self.hanging_indent
            },
            'numbering': {
                'id': self.numbering_id,
                'level': self.numbering_level
            },
            'borders': self.borders,
            'shading': self.shading,
            'based_on': self.based_on,
            'next_style': self.next_style
        }


@dataclass
class WordDocumentInfo:
    """Word文档信息"""
    filename: str
    
    # 页面设置
    page_width: Optional[float] = None
    page_height: Optional[float] = None
    margin_top: Optional[float] = None
    margin_bottom: Optional[float] = None
    margin_left: Optional[float] = None
    margin_right: Optional[float] = None
    orientation: str = 'portrait'
    
    # 页眉页脚
    header_distance: Optional[float] = None
    footer_distance: Optional[float] = None
    different_first_page: bool = False
    different_odd_even: bool = False
    
    # 样式信息
    styles: List[WordStyleInfo] = field(default_factory=list)
    
    # 编号方案
    numbering_definitions: Dict[str, Any] = field(default_factory=dict)
    
    # 主题颜色
    theme_colors: Dict[str, str] = field(default_factory=dict)
    
    # 字体方案
    font_scheme: Dict[str, str] = field(default_factory=dict)
    
    def to_dict(self) -> Dict[str, Any]:
        """转换为字典格式"""
        return {
            'filename': self.filename,
            'page_setup': {
                'width': self.page_width,
                'height': self.page_height,
                'margins': {
                    'top': self.margin_top,
                    'bottom': self.margin_bottom,
                    'left': self.margin_left,
                    'right': self.margin_right
                },
                'orientation': self.orientation
            },
            'header_footer': {
                'header_distance': self.header_distance,
                'footer_distance': self.footer_distance,
                'different_first_page': self.different_first_page,
                'different_odd_even': self.different_odd_even
            },
            'styles': [style.to_dict() for style in self.styles],
            'numbering_definitions': self.numbering_definitions,
            'theme_colors': self.theme_colors,
            'font_scheme': self.font_scheme
        }


class WordTemplateAnalyzer:
    """Word模板分析器"""
    
    def __init__(self):
        self.ns = {
            'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
            'a': 'http://schemas.openxmlformats.org/drawingml/2006/main',
            'wp': 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing',
            'r': 'http://schemas.openxmlformats.org/officeDocument/2006/relationships'
        }
        logger.info("Word模板分析器初始化完成")
    
    def analyze_word_template(self, file_path: str) -> WordDocumentInfo:
        """
        分析Word模板文件
        
        Args:
            file_path: Word文档路径
            
        Returns:
            WordDocumentInfo: 文档分析结果
        """
        logger.info(f"开始分析Word模板: {file_path}")
        
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"文件不存在: {file_path}")
        
        filename = os.path.basename(file_path)
        doc_info = WordDocumentInfo(filename=filename)
        
        try:
            # 使用python-docx分析
            doc = Document(file_path)
            self._analyze_with_docx(doc, doc_info)
            
            # 使用XML分析（更详细的信息）
            self._analyze_with_xml(file_path, doc_info)
            
            logger.info(f"模板分析完成: {len(doc_info.styles)} 个样式")
            return doc_info
            
        except Exception as e:
            logger.error(f"分析Word模板失败: {e}")
            raise
    
    def _analyze_with_docx(self, doc: Document, doc_info: WordDocumentInfo):
        """使用python-docx分析文档"""
        # 分析页面设置
        section = doc.sections[0]
        doc_info.page_width = self._twips_to_cm(section.page_width.twips if section.page_width else 11906)
        doc_info.page_height = self._twips_to_cm(section.page_height.twips if section.page_height else 16838)
        doc_info.margin_top = self._twips_to_cm(section.top_margin.twips if section.top_margin else 1440)
        doc_info.margin_bottom = self._twips_to_cm(section.bottom_margin.twips if section.bottom_margin else 1440)
        doc_info.margin_left = self._twips_to_cm(section.left_margin.twips if section.left_margin else 1440)
        doc_info.margin_right = self._twips_to_cm(section.right_margin.twips if section.right_margin else 1440)
        
        # 分析样式
        for style in doc.styles:
            style_info = self._extract_style_info(style)
            if style_info:
                doc_info.styles.append(style_info)
    
    def _analyze_with_xml(self, file_path: str, doc_info: WordDocumentInfo):
        """使用XML分析文档（获取更详细信息）"""
        try:
            with zipfile.ZipFile(file_path, 'r') as docx_zip:
                # 分析样式XML
                if 'word/styles.xml' in docx_zip.namelist():
                    styles_xml = docx_zip.read('word/styles.xml')
                    self._parse_styles_xml(styles_xml, doc_info)
                
                # 分析编号XML
                if 'word/numbering.xml' in docx_zip.namelist():
                    numbering_xml = docx_zip.read('word/numbering.xml')
                    self._parse_numbering_xml(numbering_xml, doc_info)
                
                # 分析主题XML
                if 'word/theme/theme1.xml' in docx_zip.namelist():
                    theme_xml = docx_zip.read('word/theme/theme1.xml')
                    self._parse_theme_xml(theme_xml, doc_info)
                
                # 分析文档设置XML
                if 'word/settings.xml' in docx_zip.namelist():
                    settings_xml = docx_zip.read('word/settings.xml')
                    self._parse_settings_xml(settings_xml, doc_info)
                    
        except Exception as e:
            logger.warning(f"XML分析失败: {e}")
    
    def _extract_style_info(self, style) -> Optional[WordStyleInfo]:
        """提取样式信息"""
        try:
            style_info = WordStyleInfo(
                style_id=style.style_id,
                name=style.name,
                style_type=self._get_style_type_name(style.type)
            )
            
            # 提取字体信息
            if hasattr(style, 'font'):
                font = style.font
                style_info.font_name = font.name
                if font.size:
                    style_info.font_size = font.size.pt
                style_info.font_bold = font.bold or False
                style_info.font_italic = font.italic or False
                style_info.font_underline = font.underline or False
                
                if font.color and font.color.rgb:
                    style_info.font_color = str(font.color.rgb)
            
            # 提取段落信息
            if hasattr(style, 'paragraph_format'):
                para_format = style.paragraph_format
                
                if para_format.alignment:
                    style_info.alignment = self._get_alignment_name(para_format.alignment)
                
                if para_format.space_before:
                    style_info.space_before = para_format.space_before.pt
                if para_format.space_after:
                    style_info.space_after = para_format.space_after.pt
                if para_format.line_spacing:
                    style_info.line_spacing = para_format.line_spacing
                if para_format.first_line_indent:
                    style_info.first_line_indent = self._twips_to_cm(para_format.first_line_indent.twips)
                if para_format.left_indent:
                    style_info.left_indent = self._twips_to_cm(para_format.left_indent.twips)
                if para_format.right_indent:
                    style_info.right_indent = self._twips_to_cm(para_format.right_indent.twips)
            
            # 基于的样式
            if hasattr(style, 'base_style') and style.base_style:
                style_info.based_on = style.base_style.name
            
            return style_info
            
        except Exception as e:
            logger.warning(f"提取样式 {style.name} 失败: {e}")
            return None
    
    def _parse_styles_xml(self, styles_xml: bytes, doc_info: WordDocumentInfo):
        """解析styles.xml文件"""
        try:
            root = ET.fromstring(styles_xml)
            
            for style_elem in root.findall('.//w:style', self.ns):
                style_id = style_elem.get('{%s}styleId' % self.ns['w'])
                style_type = style_elem.get('{%s}type' % self.ns['w'])
                
                # 查找对应的样式信息
                existing_style = None
                for style in doc_info.styles:
                    if style.style_id == style_id:
                        existing_style = style
                        break
                
                if existing_style:
                    # 补充XML中的详细信息
                    self._enhance_style_from_xml(style_elem, existing_style)
                    
        except Exception as e:
            logger.warning(f"解析styles.xml失败: {e}")
    
    def _parse_numbering_xml(self, numbering_xml: bytes, doc_info: WordDocumentInfo):
        """解析numbering.xml文件"""
        try:
            root = ET.fromstring(numbering_xml)
            
            # 解析抽象编号定义
            for abstractNum in root.findall('.//w:abstractNum', self.ns):
                abstractNumId = abstractNum.get('{%s}abstractNumId' % self.ns['w'])
                
                num_info = {
                    'abstractNumId': abstractNumId,
                    'levels': []
                }
                
                # 解析各级编号
                for lvl in abstractNum.findall('.//w:lvl', self.ns):
                    level_info = self._parse_numbering_level(lvl)
                    num_info['levels'].append(level_info)
                
                doc_info.numbering_definitions[abstractNumId] = num_info
                
        except Exception as e:
            logger.warning(f"解析numbering.xml失败: {e}")
    
    def _parse_theme_xml(self, theme_xml: bytes, doc_info: WordDocumentInfo):
        """解析theme.xml文件"""
        try:
            root = ET.fromstring(theme_xml)
            
            # 解析主题颜色
            color_scheme = root.find('.//a:clrScheme', self.ns)
            if color_scheme:
                for color_elem in color_scheme:
                    color_name = color_elem.tag.split('}')[-1]  # 去掉命名空间
                    
                    # 查找颜色值
                    srgb_clr = color_elem.find('.//a:srgbClr', self.ns)
                    if srgb_clr is not None:
                        color_val = srgb_clr.get('val')
                        doc_info.theme_colors[color_name] = color_val
            
            # 解析字体方案
            font_scheme = root.find('.//a:fontScheme', self.ns)
            if font_scheme:
                major_font = font_scheme.find('.//a:majorFont/a:latin', self.ns)
                minor_font = font_scheme.find('.//a:minorFont/a:latin', self.ns)
                
                if major_font is not None:
                    doc_info.font_scheme['major'] = major_font.get('typeface')
                if minor_font is not None:
                    doc_info.font_scheme['minor'] = minor_font.get('typeface')
                    
        except Exception as e:
            logger.warning(f"解析theme.xml失败: {e}")
    
    def _parse_settings_xml(self, settings_xml: bytes, doc_info: WordDocumentInfo):
        """解析settings.xml文件"""
        try:
            root = ET.fromstring(settings_xml)
            
            # 查找特殊设置
            even_and_odd_headers = root.find('.//w:evenAndOddHeaders', self.ns)
            if even_and_odd_headers is not None:
                doc_info.different_odd_even = True
                
        except Exception as e:
            logger.warning(f"解析settings.xml失败: {e}")
    
    def _enhance_style_from_xml(self, style_elem, style_info: WordStyleInfo):
        """从XML元素中增强样式信息"""
        try:
            # 解析边框信息
            pBdr = style_elem.find('.//w:pBdr', self.ns)
            if pBdr is not None:
                borders = {}
                for border_type in ['top', 'left', 'bottom', 'right']:
                    border = pBdr.find(f'.//w:{border_type}', self.ns)
                    if border is not None:
                        borders[border_type] = {
                            'val': border.get('{%s}val' % self.ns['w']),
                            'sz': border.get('{%s}sz' % self.ns['w']),
                            'color': border.get('{%s}color' % self.ns['w'])
                        }
                style_info.borders = borders
            
            # 解析底纹信息
            shd = style_elem.find('.//w:shd', self.ns)
            if shd is not None:
                style_info.shading = {
                    'val': shd.get('{%s}val' % self.ns['w']),
                    'color': shd.get('{%s}color' % self.ns['w']),
                    'fill': shd.get('{%s}fill' % self.ns['w'])
                }
            
            # 解析编号信息
            numPr = style_elem.find('.//w:numPr', self.ns)
            if numPr is not None:
                numId = numPr.find('.//w:numId', self.ns)
                ilvl = numPr.find('.//w:ilvl', self.ns)
                
                if numId is not None:
                    style_info.numbering_id = numId.get('{%s}val' % self.ns['w'])
                if ilvl is not None:
                    style_info.numbering_level = int(ilvl.get('{%s}val' % self.ns['w']))
                    
        except Exception as e:
            logger.warning(f"增强样式信息失败: {e}")
    
    def _parse_numbering_level(self, lvl_elem) -> Dict[str, Any]:
        """解析编号级别信息"""
        level_info = {
            'ilvl': lvl_elem.get('{%s}ilvl' % self.ns['w']),
            'start': None,
            'numFmt': None,
            'lvlText': None,
            'lvlJc': None
        }
        
        try:
            start = lvl_elem.find('.//w:start', self.ns)
            if start is not None:
                level_info['start'] = start.get('{%s}val' % self.ns['w'])
            
            numFmt = lvl_elem.find('.//w:numFmt', self.ns)
            if numFmt is not None:
                level_info['numFmt'] = numFmt.get('{%s}val' % self.ns['w'])
            
            lvlText = lvl_elem.find('.//w:lvlText', self.ns)
            if lvlText is not None:
                level_info['lvlText'] = lvlText.get('{%s}val' % self.ns['w'])
            
            lvlJc = lvl_elem.find('.//w:lvlJc', self.ns)
            if lvlJc is not None:
                level_info['lvlJc'] = lvlJc.get('{%s}val' % self.ns['w'])
                
        except Exception as e:
            logger.warning(f"解析编号级别失败: {e}")
        
        return level_info
    
    def _get_style_type_name(self, style_type) -> str:
        """获取样式类型名称"""
        type_map = {
            WD_STYLE_TYPE.PARAGRAPH: 'paragraph',
            WD_STYLE_TYPE.CHARACTER: 'character',
            WD_STYLE_TYPE.TABLE: 'table',
            WD_STYLE_TYPE.LIST: 'list'
        }
        return type_map.get(style_type, 'unknown')
    
    def _get_alignment_name(self, alignment) -> str:
        """获取对齐方式名称"""
        alignment_map = {
            WD_PARAGRAPH_ALIGNMENT.LEFT: 'left',
            WD_PARAGRAPH_ALIGNMENT.CENTER: 'center',
            WD_PARAGRAPH_ALIGNMENT.RIGHT: 'right',
            WD_PARAGRAPH_ALIGNMENT.JUSTIFY: 'justify',
            WD_PARAGRAPH_ALIGNMENT.DISTRIBUTE: 'distribute'
        }
        return alignment_map.get(alignment, 'left')
    
    def _twips_to_cm(self, twips: int) -> float:
        """将twips转换为厘米"""
        return twips / 566.929  # 1 cm = 566.929 twips
    
    def _cm_to_twips(self, cm: float) -> int:
        """将厘米转换为twips"""
        return int(cm * 566.929)
    
    def export_template_config(self, doc_info: WordDocumentInfo, output_path: str):
        """导出模板配置为JSON文件"""
        try:
            config_data = doc_info.to_dict()
            
            with open(output_path, 'w', encoding='utf-8') as f:
                json.dump(config_data, f, ensure_ascii=False, indent=2)
            
            logger.info(f"模板配置已导出到: {output_path}")
            
        except Exception as e:
            logger.error(f"导出模板配置失败: {e}")
            raise
    
    def compare_templates(self, template1: WordDocumentInfo, 
                         template2: WordDocumentInfo) -> Dict[str, Any]:
        """比较两个模板的差异"""
        comparison = {
            'page_setup_diff': {},
            'styles_diff': {
                'only_in_template1': [],
                'only_in_template2': [],
                'different_styles': []
            },
            'numbering_diff': {},
            'theme_diff': {}
        }
        
        try:
            # 比较页面设置
            if template1.page_width != template2.page_width:
                comparison['page_setup_diff']['page_width'] = {
                    'template1': template1.page_width,
                    'template2': template2.page_width
                }
            
            # 比较样式
            styles1 = {s.style_id: s for s in template1.styles}
            styles2 = {s.style_id: s for s in template2.styles}
            
            # 仅在模板1中的样式
            comparison['styles_diff']['only_in_template1'] = [
                s.name for s_id, s in styles1.items() if s_id not in styles2
            ]
            
            # 仅在模板2中的样式
            comparison['styles_diff']['only_in_template2'] = [
                s.name for s_id, s in styles2.items() if s_id not in styles1
            ]
            
            # 不同的样式
            for s_id in set(styles1.keys()) & set(styles2.keys()):
                if self._styles_different(styles1[s_id], styles2[s_id]):
                    comparison['styles_diff']['different_styles'].append({
                        'style_name': styles1[s_id].name,
                        'differences': self._get_style_differences(styles1[s_id], styles2[s_id])
                    })
            
            return comparison
            
        except Exception as e:
            logger.error(f"比较模板失败: {e}")
            return comparison
    
    def _styles_different(self, style1: WordStyleInfo, style2: WordStyleInfo) -> bool:
        """检查两个样式是否不同"""
        # 简化的比较逻辑
        return (style1.font_name != style2.font_name or
                style1.font_size != style2.font_size or
                style1.alignment != style2.alignment)
    
    def _get_style_differences(self, style1: WordStyleInfo, 
                              style2: WordStyleInfo) -> Dict[str, Any]:
        """获取样式差异详情"""
        differences = {}
        
        if style1.font_name != style2.font_name:
            differences['font_name'] = {
                'style1': style1.font_name,
                'style2': style2.font_name
            }
        
        if style1.font_size != style2.font_size:
            differences['font_size'] = {
                'style1': style1.font_size,
                'style2': style2.font_size
            }
        
        if style1.alignment != style2.alignment:
            differences['alignment'] = {
                'style1': style1.alignment,
                'style2': style2.alignment
            }
        
        return differences
    
    def extract_content_structure(self, file_path: str) -> Dict[str, Any]:
        """提取文档内容结构（用于理解模板的内容组织）"""
        try:
            doc = Document(file_path)
            structure = {
                'paragraphs': [],
                'tables': [],
                'headers': [],
                'footers': []
            }
            
            # 分析段落结构
            for i, para in enumerate(doc.paragraphs):
                para_info = {
                    'index': i,
                    'text': para.text[:100],  # 前100字符
                    'style': para.style.name if para.style else 'Normal',
                    'level': self._get_paragraph_level(para),
                    'is_heading': self._is_heading_paragraph(para)
                }
                structure['paragraphs'].append(para_info)
            
            # 分析表格
            for i, table in enumerate(doc.tables):
                table_info = {
                    'index': i,
                    'rows': len(table.rows),
                    'cols': len(table.columns) if table.rows else 0,
                    'style': table.style.name if table.style else None
                }
                structure['tables'].append(table_info)
            
            # 分析页眉页脚
            for section in doc.sections:
                if section.header:
                    header_info = {
                        'text': section.header.paragraphs[0].text if section.header.paragraphs else '',
                        'is_linked': section.header.is_linked_to_previous
                    }
                    structure['headers'].append(header_info)
                
                if section.footer:
                    footer_info = {
                        'text': section.footer.paragraphs[0].text if section.footer.paragraphs else '',
                        'is_linked': section.footer.is_linked_to_previous
                    }
                    structure['footers'].append(footer_info)
            
            return structure
            
        except Exception as e:
            logger.error(f"提取内容结构失败: {e}")
            return {}
    
    def _get_paragraph_level(self, paragraph) -> int:
        """获取段落级别"""
        style_name = paragraph.style.name.lower()
        
        # 检查是否是标题样式
        if 'heading' in style_name:
            try:
                level = int(re.search(r'heading\s*(\d+)', style_name).group(1))
                return level
            except:
                pass
        
        # 检查编号级别
        if paragraph._p.pPr is not None:
            numPr = paragraph._p.pPr.find(qn('w:numPr'))
            if numPr is not None:
                ilvl = numPr.find(qn('w:ilvl'))
                if ilvl is not None:
                    return int(ilvl.get(qn('w:val'))) + 1
        
        return 0
    
    def _is_heading_paragraph(self, paragraph) -> bool:
        """判断是否是标题段落"""
        style_name = paragraph.style.name.lower()
        return 'heading' in style_name or 'title' in style_name


class TemplateLibrary:
    """模板库管理器"""
    
    def __init__(self, library_path: str = "template_library"):
        self.library_path = Path(library_path)
        self.library_path.mkdir(exist_ok=True)
        self.analyzer = WordTemplateAnalyzer()
        
        # 模板索引文件
        self.index_file = self.library_path / "template_index.json"
        self.template_index = self._load_template_index()
    
    def _load_template_index(self) -> Dict[str, Any]:
        """加载模板索引"""
        if self.index_file.exists():
            try:
                with open(self.index_file, 'r', encoding='utf-8') as f:
                    return json.load(f)
            except Exception as e:
                logger.warning(f"加载模板索引失败: {e}")
        
        return {"templates": {}, "version": "1.0"}
    
    def _save_template_index(self):
        """保存模板索引"""
        try:
            with open(self.index_file, 'w', encoding='utf-8') as f:
                json.dump(self.template_index, f, ensure_ascii=False, indent=2)
        except Exception as e:
            logger.error(f"保存模板索引失败: {e}")
    
    def add_template(self, word_file_path: str, template_name: str, 
                    description: str = "", tags: List[str] = None) -> str:
        """
        添加Word模板到库中
        
        Args:
            word_file_path: Word文件路径
            template_name: 模板名称
            description: 模板描述
            tags: 标签列表
            
        Returns:
            str: 模板ID
        """
        try:
            # 分析Word模板
            doc_info = self.analyzer.analyze_word_template(word_file_path)
            
            # 生成模板ID
            template_id = self._generate_template_id(template_name)
            
            # 创建模板目录
            template_dir = self.library_path / template_id
            template_dir.mkdir(exist_ok=True)
            
            # 复制原始Word文件
            import shutil
            word_copy_path = template_dir / f"{template_id}.docx"
            shutil.copy2(word_file_path, word_copy_path)
            
            # 保存分析结果
            config_path = template_dir / "template_config.json"
            self.analyzer.export_template_config(doc_info, str(config_path))
            
            # 提取内容结构
            structure = self.analyzer.extract_content_structure(word_file_path)
            structure_path = template_dir / "content_structure.json"
            with open(structure_path, 'w', encoding='utf-8') as f:
                json.dump(structure, f, ensure_ascii=False, indent=2)
            
            # 更新索引
            self.template_index["templates"][template_id] = {
                "name": template_name,
                "description": description,
                "tags": tags or [],
                "created_at": self._get_current_timestamp(),
                "word_file": str(word_copy_path),
                "config_file": str(config_path),
                "structure_file": str(structure_path),
                "styles_count": len(doc_info.styles),
                "page_setup": {
                    "width": doc_info.page_width,
                    "height": doc_info.page_height,
                    "orientation": doc_info.orientation
                }
            }
            
            self._save_template_index()
            
            logger.info(f"模板已添加到库中: {template_id}")
            return template_id
            
        except Exception as e:
            logger.error(f"添加模板失败: {e}")
            raise
    
    def get_template_info(self, template_id: str) -> Optional[WordDocumentInfo]:
        """获取模板信息"""
        if template_id not in self.template_index["templates"]:
            return None
        
        try:
            template_info = self.template_index["templates"][template_id]
            config_file = template_info["config_file"]
            
            with open(config_file, 'r', encoding='utf-8') as f:
                config_data = json.load(f)
            
            # 重构WordDocumentInfo对象
            doc_info = WordDocumentInfo(filename=config_data["filename"])
            doc_info.page_width = config_data["page_setup"]["width"]
            doc_info.page_height = config_data["page_setup"]["height"]
            doc_info.orientation = config_data["page_setup"]["orientation"]
            
            # 重构样式信息
            for style_data in config_data["styles"]:
                style_info = WordStyleInfo(
                    style_id=style_data["style_id"],
                    name=style_data["name"],
                    style_type=style_data["style_type"]
                )
                # 填充其他属性...
                doc_info.styles.append(style_info)
            
            return doc_info
            
        except Exception as e:
            logger.error(f"获取模板信息失败: {e}")
            return None
    
    def list_templates(self) -> Dict[str, Any]:
        """列出所有模板"""
        return self.template_index["templates"]
    
    def remove_template(self, template_id: str) -> bool:
        """删除模板"""
        if template_id not in self.template_index["templates"]:
            return False
        
        try:
            # 删除模板目录
            template_dir = self.library_path / template_id
            if template_dir.exists():
                import shutil
                shutil.rmtree(template_dir)
            
            # 从索引中删除
            del self.template_index["templates"][template_id]
            self._save_template_index()
            
            logger.info(f"模板已删除: {template_id}")
            return True
            
        except Exception as e:
            logger.error(f"删除模板失败: {e}")
            return False
    
    def search_templates(self, query: str = "", tags: List[str] = None) -> Dict[str, Any]:
        """搜索模板"""
        results = {}
        
        for template_id, template_info in self.template_index["templates"].items():
            # 文本搜索
            if query:
                if (query.lower() in template_info["name"].lower() or
                    query.lower() in template_info["description"].lower()):
                    results[template_id] = template_info
                    continue
            
            # 标签搜索
            if tags:
                template_tags = template_info.get("tags", [])
                if any(tag in template_tags for tag in tags):
                    results[template_id] = template_info
                    continue
            
            # 如果没有搜索条件，返回所有
            if not query and not tags:
                results[template_id] = template_info
        
        return results
    
    def _generate_template_id(self, template_name: str) -> str:
        """生成模板ID"""
        import hashlib
        import time
        
        # 使用名称和时间戳生成唯一ID
        content = f"{template_name}_{time.time()}"
        return hashlib.md5(content.encode()).hexdigest()[:12]
    
    def _get_current_timestamp(self) -> str:
        """获取当前时间戳"""
        import datetime
        return datetime.datetime.now().isoformat()


# 便捷函数
def analyze_word_template(file_path: str) -> WordDocumentInfo:
    """分析Word模板的便捷函数"""
    analyzer = WordTemplateAnalyzer()
    return analyzer.analyze_word_template(file_path)


def create_template_library(library_path: str = "template_library") -> TemplateLibrary:
    """创建模板库的便捷函数"""
    return TemplateLibrary(library_path)