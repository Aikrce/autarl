#!/usr/bin/env python3
"""
Template-Based Conversion Engine
基于模板的精确转换引擎 - 实现从Markdown到Word的高精度样式转换
"""

import os
import re
import logging
from typing import Dict, List, Optional, Any, Tuple, Union
from dataclasses import dataclass
from pathlib import Path
import tempfile
import shutil

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Length
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import markdown2

from word_template_analyzer import WordDocumentInfo, WordStyleInfo, analyze_word_template
from markdown_style_mapper import (
    MarkdownWordStyleMapper, MarkdownElementType, StyleMapping,
    MarkdownAnalyzer, analyze_markdown_for_mapping
)
from enhanced_document_analyzer import analyze_markdown_document, ContentSection, SectionType

logger = logging.getLogger(__name__)


@dataclass
class ConversionResult:
    """转换结果"""
    success: bool
    output_path: str
    template_used: str
    style_mappings_applied: int
    warnings: List[str]
    errors: List[str]
    processing_time: float
    
    def to_dict(self) -> Dict[str, Any]:
        return {
            'success': self.success,
            'output_path': self.output_path,
            'template_used': self.template_used,
            'style_mappings_applied': self.style_mappings_applied,
            'warnings': self.warnings,
            'errors': self.errors,
            'processing_time': self.processing_time
        }


class TemplateBasedConverter:
    """基于模板的转换器"""
    
    def __init__(self, word_template_path: str):
        """
        初始化转换器
        
        Args:
            word_template_path: Word模板文件路径
        """
        self.template_path = word_template_path
        self.template_info = analyze_word_template(word_template_path)
        self.style_mapper = MarkdownWordStyleMapper(self.template_info)
        self.markdown_analyzer = MarkdownAnalyzer()
        
        # 转换统计
        self.conversion_stats = {
            'total_conversions': 0,
            'successful_conversions': 0,
            'failed_conversions': 0,
            'total_style_applications': 0
        }
        
        logger.info(f"模板转换器初始化完成: {self.template_info.filename}")
    
    def convert_markdown_to_word(self, markdown_content: str, output_path: str,
                                preserve_template_structure: bool = True) -> ConversionResult:
        """
        将Markdown转换为Word文档
        
        Args:
            markdown_content: Markdown内容
            output_path: 输出Word文档路径
            preserve_template_structure: 是否保留模板结构
            
        Returns:
            ConversionResult: 转换结果
        """
        import time
        start_time = time.time()
        
        warnings = []
        errors = []
        style_mappings_applied = 0
        
        try:
            self.conversion_stats['total_conversions'] += 1
            
            # 分析Markdown文档
            logger.info("分析Markdown文档结构...")
            document_structure = analyze_markdown_document(markdown_content)
            document_context = self.markdown_analyzer.analyze_document_context(markdown_content)
            
            # 创建新的Word文档，基于模板
            logger.info("基于模板创建Word文档...")
            doc = self._create_document_from_template()
            
            # 设置页面格式
            self._apply_page_setup(doc)
            
            # 解析Markdown内容
            parsed_sections = self._parse_markdown_content(markdown_content)
            
            # 按照文档结构应用内容
            if preserve_template_structure and document_structure.sections:
                style_mappings_applied = self._apply_structured_content(
                    doc, document_structure.sections, document_context, warnings, errors
                )
            else:
                style_mappings_applied = self._apply_sequential_content(
                    doc, parsed_sections, document_context, warnings, errors
                )
            
            # 应用页眉页脚
            self._apply_headers_footers(doc)
            
            # 保存文档
            doc.save(output_path)
            
            # 更新统计
            self.conversion_stats['successful_conversions'] += 1
            self.conversion_stats['total_style_applications'] += style_mappings_applied
            
            processing_time = time.time() - start_time
            
            logger.info(f"转换完成: {output_path}, 应用了 {style_mappings_applied} 个样式映射")
            
            return ConversionResult(
                success=True,
                output_path=output_path,
                template_used=self.template_info.filename,
                style_mappings_applied=style_mappings_applied,
                warnings=warnings,
                errors=errors,
                processing_time=processing_time
            )
            
        except Exception as e:
            self.conversion_stats['failed_conversions'] += 1
            processing_time = time.time() - start_time
            
            error_msg = f"转换失败: {str(e)}"
            logger.error(error_msg)
            errors.append(error_msg)
            
            return ConversionResult(
                success=False,
                output_path="",
                template_used=self.template_info.filename,
                style_mappings_applied=0,
                warnings=warnings,
                errors=errors,
                processing_time=processing_time
            )
    
    def _create_document_from_template(self) -> Document:
        """基于模板创建Word文档"""
        try:
            # 创建基于模板的文档
            doc = Document(self.template_path)
            
            # 清除模板中的示例内容，但保留样式
            self._clear_template_content(doc)
            
            return doc
            
        except Exception as e:
            logger.warning(f"无法基于模板创建文档，创建空白文档: {e}")
            # 创建空白文档并应用模板样式
            doc = Document()
            self._apply_template_styles(doc)
            return doc
    
    def _clear_template_content(self, doc: Document):
        """清除模板内容但保留样式"""
        # 清除所有段落但保留第一个（用于保持样式）
        paragraphs_to_remove = doc.paragraphs[1:]  # 保留第一个段落
        
        for paragraph in paragraphs_to_remove:
            p = paragraph._element
            p.getparent().remove(p)
        
        # 清空第一个段落的文本
        if doc.paragraphs:
            doc.paragraphs[0].clear()
        
        # 清除表格
        for table in doc.tables:
            t = table._element
            t.getparent().remove(t)
    
    def _apply_template_styles(self, doc: Document):
        """将模板样式应用到新文档"""
        # 这里可以实现从模板信息重建样式的逻辑
        # 由于python-docx的限制，这个功能比较复杂
        logger.info("应用模板样式到新文档")
    
    def _apply_page_setup(self, doc: Document):
        """应用页面设置"""
        try:
            section = doc.sections[0]
            
            # 应用页面尺寸
            if self.template_info.page_width:
                section.page_width = Cm(self.template_info.page_width)
            if self.template_info.page_height:
                section.page_height = Cm(self.template_info.page_height)
            
            # 应用页边距
            if self.template_info.margin_top:
                section.top_margin = Cm(self.template_info.margin_top)
            if self.template_info.margin_bottom:
                section.bottom_margin = Cm(self.template_info.margin_bottom)
            if self.template_info.margin_left:
                section.left_margin = Cm(self.template_info.margin_left)
            if self.template_info.margin_right:
                section.right_margin = Cm(self.template_info.margin_right)
            
            logger.debug("页面设置已应用")
            
        except Exception as e:
            logger.warning(f"应用页面设置失败: {e}")
    
    def _parse_markdown_content(self, content: str) -> List[Dict[str, Any]]:
        """解析Markdown内容为结构化数据"""
        parsed_sections = []
        lines = content.split('\n')
        
        current_section = None
        current_content = []
        
        for line in lines:
            line_stripped = line.strip()
            
            # 检测标题
            heading_match = re.match(r'^(#{1,6})\s+(.+)$', line_stripped)
            if heading_match:
                # 保存之前的段落
                if current_section:
                    current_section['content'] = '\n'.join(current_content)
                    parsed_sections.append(current_section)
                
                # 开始新段落
                level = len(heading_match.group(1))
                title = heading_match.group(2).strip()
                
                current_section = {
                    'type': f'heading_{level}',
                    'element_type': getattr(MarkdownElementType, f'HEADING_{level}'),
                    'title': title,
                    'level': level,
                    'raw_line': line
                }
                current_content = []
            
            else:
                if current_section is None:
                    # 创建默认段落
                    current_section = {
                        'type': 'paragraph',
                        'element_type': MarkdownElementType.PARAGRAPH,
                        'title': '',
                        'level': 0,
                        'raw_line': ''
                    }
                    current_content = []
                
                current_content.append(line)
        
        # 保存最后一个段落
        if current_section:
            current_section['content'] = '\n'.join(current_content)
            parsed_sections.append(current_section)
        
        return parsed_sections
    
    def _apply_structured_content(self, doc: Document, sections: List[ContentSection],
                                 document_context: Dict[str, Any],
                                 warnings: List[str], errors: List[str]) -> int:
        """应用结构化内容"""
        style_count = 0
        
        for section in sections:
            try:
                # 确定元素类型
                element_type = self._map_section_to_element_type(section)
                
                # 获取样式映射
                style_mapping = self.style_mapper.get_contextual_style(
                    section.name, element_type, document_context
                )
                
                if not style_mapping:
                    style_mapping = self.style_mapper.get_style_mapping(element_type)
                
                # 应用内容
                if section.section_type in [SectionType.ABSTRACT_CN, SectionType.ABSTRACT_EN]:
                    self._add_abstract_section(doc, section, style_mapping, document_context)
                elif section.section_type == SectionType.REFERENCES:
                    self._add_references_section(doc, section, style_mapping)
                elif section.section_type == SectionType.TOC:
                    self._add_toc_section(doc, section, style_mapping)
                else:
                    self._add_regular_section(doc, section, style_mapping)
                
                if style_mapping:
                    style_count += 1
                    
            except Exception as e:
                error_msg = f"处理章节 '{section.name}' 失败: {str(e)}"
                logger.warning(error_msg)
                warnings.append(error_msg)
        
        return style_count
    
    def _apply_sequential_content(self, doc: Document, parsed_sections: List[Dict[str, Any]],
                                 document_context: Dict[str, Any],
                                 warnings: List[str], errors: List[str]) -> int:
        """按顺序应用内容"""
        style_count = 0
        
        for section_data in parsed_sections:
            try:
                element_type = section_data['element_type']
                
                # 获取样式映射
                if section_data.get('title'):
                    style_mapping = self.style_mapper.get_contextual_style(
                        section_data['title'], element_type, document_context
                    )
                else:
                    style_mapping = self.style_mapper.get_style_mapping(element_type)
                
                # 添加内容
                if element_type in [MarkdownElementType.HEADING_1, MarkdownElementType.HEADING_2,
                                   MarkdownElementType.HEADING_3, MarkdownElementType.HEADING_4,
                                   MarkdownElementType.HEADING_5, MarkdownElementType.HEADING_6]:
                    self._add_heading(doc, section_data['title'], style_mapping)
                else:
                    self._add_paragraph_content(doc, section_data['content'], style_mapping)
                
                if style_mapping:
                    style_count += 1
                    
            except Exception as e:
                error_msg = f"处理内容段落失败: {str(e)}"
                logger.warning(error_msg)
                warnings.append(error_msg)
        
        return style_count
    
    def _map_section_to_element_type(self, section: ContentSection) -> MarkdownElementType:
        """将章节类型映射到Markdown元素类型"""
        mapping = {
            SectionType.ABSTRACT_CN: MarkdownElementType.ABSTRACT_TITLE,
            SectionType.ABSTRACT_EN: MarkdownElementType.ABSTRACT_TITLE,
            SectionType.KEYWORDS_CN: MarkdownElementType.KEYWORDS,
            SectionType.KEYWORDS_EN: MarkdownElementType.KEYWORDS,
            SectionType.TOC: MarkdownElementType.TOC_TITLE,
            SectionType.REFERENCES: MarkdownElementType.REFERENCE_TITLE,
            SectionType.CHAPTER: MarkdownElementType.CHAPTER_TITLE,
            SectionType.SECTION: MarkdownElementType.SECTION_TITLE,
            SectionType.SUBSECTION: MarkdownElementType.HEADING_3
        }
        
        return mapping.get(section.section_type, MarkdownElementType.PARAGRAPH)
    
    def _add_abstract_section(self, doc: Document, section: ContentSection,
                             style_mapping: Optional[StyleMapping],
                             document_context: Dict[str, Any]):
        """添加摘要章节"""
        # 添加分页符（摘要通常独立成页）
        doc.add_page_break()
        
        # 添加摘要标题
        title_para = doc.add_paragraph(section.name)
        if style_mapping:
            self._apply_style_to_paragraph(title_para, style_mapping.word_style_name)
        
        # 处理摘要内容
        content_lines = section.content.split('\n')[1:]  # 跳过标题行
        
        # 查找摘要正文样式
        content_style_mapping = self.style_mapper.get_style_mapping(MarkdownElementType.ABSTRACT_CONTENT)
        
        # 添加摘要正文
        for line in content_lines:
            line = line.strip()
            if line and not line.startswith('关键词') and not line.startswith('Key words'):
                para = doc.add_paragraph(line)
                if content_style_mapping:
                    self._apply_style_to_paragraph(para, content_style_mapping.word_style_name)
        
        # 处理关键词
        keywords_line = self._extract_keywords_line(section.content)
        if keywords_line:
            keywords_para = doc.add_paragraph(keywords_line)
            keywords_style_mapping = self.style_mapper.get_style_mapping(MarkdownElementType.KEYWORDS)
            if keywords_style_mapping:
                self._apply_style_to_paragraph(keywords_para, keywords_style_mapping.word_style_name)
    
    def _add_references_section(self, doc: Document, section: ContentSection,
                               style_mapping: Optional[StyleMapping]):
        """添加参考文献章节"""
        # 添加参考文献标题
        title_para = doc.add_paragraph(section.name)
        if style_mapping:
            self._apply_style_to_paragraph(title_para, style_mapping.word_style_name)
        
        # 处理参考文献条目
        content_lines = section.content.split('\n')[1:]  # 跳过标题行
        ref_item_mapping = self.style_mapper.get_style_mapping(MarkdownElementType.REFERENCE_ITEM)
        
        for line in content_lines:
            line = line.strip()
            if line:
                ref_para = doc.add_paragraph(line)
                if ref_item_mapping:
                    self._apply_style_to_paragraph(ref_para, ref_item_mapping.word_style_name)
                
                # 应用智能悬挂缩进
                self._apply_reference_hanging_indent(ref_para, line)
    
    def _add_toc_section(self, doc: Document, section: ContentSection,
                        style_mapping: Optional[StyleMapping]):
        """添加目录章节"""
        # 添加分页符
        doc.add_page_break()
        
        # 添加目录标题
        title_para = doc.add_paragraph(section.name)
        if style_mapping:
            self._apply_style_to_paragraph(title_para, style_mapping.word_style_name)
        
        # 处理目录内容
        content_lines = section.content.split('\n')[1:]  # 跳过标题行
        
        for line in content_lines:
            line = line.strip()
            if line:
                # 判断目录级别
                toc_level = self._determine_toc_level(line)
                toc_para = doc.add_paragraph(line)
                
                # 根据级别应用不同的缩进
                if toc_level == 1:
                    toc_para.paragraph_format.left_indent = Cm(0)
                elif toc_level == 2:
                    toc_para.paragraph_format.left_indent = Cm(0.37)
                elif toc_level >= 3:
                    toc_para.paragraph_format.left_indent = Cm(0.74)
    
    def _add_regular_section(self, doc: Document, section: ContentSection,
                            style_mapping: Optional[StyleMapping]):
        """添加常规章节"""
        # 为主要章节添加分页符
        if section.level == 1:
            doc.add_page_break()
        
        # 添加标题
        title_para = doc.add_paragraph(section.name)
        if style_mapping:
            self._apply_style_to_paragraph(title_para, style_mapping.word_style_name)
        
        # 添加内容
        self._add_markdown_content(doc, section.content)
    
    def _add_heading(self, doc: Document, title: str, style_mapping: Optional[StyleMapping]):
        """添加标题"""
        heading_para = doc.add_paragraph(title)
        if style_mapping:
            self._apply_style_to_paragraph(heading_para, style_mapping.word_style_name)
    
    def _add_paragraph_content(self, doc: Document, content: str,
                              style_mapping: Optional[StyleMapping]):
        """添加段落内容"""
        if not content.strip():
            return
        
        # 处理Markdown格式
        self._add_markdown_content(doc, content)
    
    def _add_markdown_content(self, doc: Document, content: str):
        """添加Markdown格式化内容"""
        lines = content.split('\n')
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
            
            # 检测特殊格式
            if line.startswith('```'):
                # 代码块开始/结束
                continue
            elif line.startswith(('- ', '* ', '+ ')):
                # 无序列表
                list_text = line[2:].strip()
                list_para = doc.add_paragraph(list_text)
                list_mapping = self.style_mapper.get_style_mapping(MarkdownElementType.LIST_UNORDERED)
                if list_mapping:
                    self._apply_style_to_paragraph(list_para, list_mapping.word_style_name)
            elif re.match(r'^\d+\.\s+', line):
                # 有序列表
                list_text = re.sub(r'^\d+\.\s+', '', line)
                list_para = doc.add_paragraph(list_text)
                list_mapping = self.style_mapper.get_style_mapping(MarkdownElementType.LIST_ORDERED)
                if list_mapping:
                    self._apply_style_to_paragraph(list_para, list_mapping.word_style_name)
            elif line.startswith('>'):
                # 引用
                quote_text = line[1:].strip()
                quote_para = doc.add_paragraph(quote_text)
                quote_mapping = self.style_mapper.get_style_mapping(MarkdownElementType.QUOTE)
                if quote_mapping:
                    self._apply_style_to_paragraph(quote_para, quote_mapping.word_style_name)
            else:
                # 普通段落
                para = doc.add_paragraph()
                self._process_inline_formatting(para, line)
    
    def _process_inline_formatting(self, paragraph, text: str):
        """处理行内格式化"""
        # 简化的行内格式处理
        parts = re.split(r'(\*\*.*?\*\*|__.*?__|_.*?_|\*.*?\*|`.*?`)', text)
        
        for part in parts:
            if not part:
                continue
            
            # 粗体
            if part.startswith('**') and part.endswith('**'):
                run = paragraph.add_run(part[2:-2])
                run.bold = True
            elif part.startswith('__') and part.endswith('__'):
                run = paragraph.add_run(part[2:-2])
                run.bold = True
            # 斜体
            elif part.startswith('*') and part.endswith('*') and len(part) > 2:
                run = paragraph.add_run(part[1:-1])
                run.italic = True
            elif part.startswith('_') and part.endswith('_') and len(part) > 2:
                run = paragraph.add_run(part[1:-1])
                run.italic = True
            # 行内代码
            elif part.startswith('`') and part.endswith('`'):
                run = paragraph.add_run(part[1:-1])
                run.font.name = 'Consolas'
                run.font.size = Pt(10)
            else:
                paragraph.add_run(part)
    
    def _apply_style_to_paragraph(self, paragraph, style_name: str):
        """应用样式到段落"""
        try:
            paragraph.style = style_name
        except Exception as e:
            logger.warning(f"应用样式 '{style_name}' 失败: {e}")
    
    def _apply_headers_footers(self, doc: Document):
        """应用页眉页脚"""
        try:
            for section in doc.sections:
                # 应用页眉
                if self.template_info.header_distance is not None:
                    section.header_distance = Cm(self.template_info.header_distance)
                
                # 应用页脚
                if self.template_info.footer_distance is not None:
                    section.footer_distance = Cm(self.template_info.footer_distance)
                
                # 设置页眉内容（如果模板有定义）
                # 这里可以根据模板信息设置页眉页脚内容
                
        except Exception as e:
            logger.warning(f"应用页眉页脚失败: {e}")
    
    def _extract_keywords_line(self, content: str) -> Optional[str]:
        """提取关键词行"""
        lines = content.split('\n')
        for line in lines:
            if '关键词' in line or 'Key words' in line:
                return line.strip()
        return None
    
    def _apply_reference_hanging_indent(self, paragraph, ref_text: str):
        """应用参考文献悬挂缩进"""
        # 检测编号格式
        number_match = re.match(r'^\[(\d+)\]', ref_text.strip())
        if number_match:
            number = int(number_match.group(1))
            if number <= 9:
                paragraph.paragraph_format.hanging_indent = Cm(0.6)
            elif number <= 99:
                paragraph.paragraph_format.hanging_indent = Cm(0.74)
            else:
                paragraph.paragraph_format.hanging_indent = Cm(0.9)
    
    def _determine_toc_level(self, line: str) -> int:
        """确定目录项级别"""
        # 清理Markdown标记
        cleaned_line = line.replace('*', '').replace('_', '').strip()
        
        # 检测章节标记
        if ('第' in cleaned_line and '章' in cleaned_line) or cleaned_line in ['绪论', '结论', '引言']:
            return 1
        elif re.match(r'^\d+\.\d+', cleaned_line):
            return 2
        elif re.match(r'^\d+\.\d+\.\d+', cleaned_line):
            return 3
        else:
            return 1
    
    def get_conversion_statistics(self) -> Dict[str, Any]:
        """获取转换统计"""
        stats = self.conversion_stats.copy()
        
        if stats['total_conversions'] > 0:
            stats['success_rate'] = stats['successful_conversions'] / stats['total_conversions'] * 100
            stats['avg_style_applications'] = stats['total_style_applications'] / stats['successful_conversions'] if stats['successful_conversions'] > 0 else 0
        else:
            stats['success_rate'] = 0.0
            stats['avg_style_applications'] = 0.0
        
        return stats
    
    def get_style_mapping_info(self) -> Dict[str, Any]:
        """获取样式映射信息"""
        return self.style_mapper.get_mapping_statistics()
    
    def export_conversion_config(self, output_path: str):
        """导出转换配置"""
        config_data = {
            'template_info': self.template_info.to_dict(),
            'style_mappings': self.style_mapper.get_mapping_statistics(),
            'conversion_stats': self.get_conversion_statistics()
        }
        
        import json
        with open(output_path, 'w', encoding='utf-8') as f:
            json.dump(config_data, f, ensure_ascii=False, indent=2)
        
        logger.info(f"转换配置已导出到: {output_path}")


class AdvancedTemplateConverter:
    """高级模板转换器 - 支持多模板和批量处理"""
    
    def __init__(self, template_library_path: str = "template_library"):
        """
        初始化高级转换器
        
        Args:
            template_library_path: 模板库路径
        """
        from word_template_analyzer import TemplateLibrary
        
        self.template_library = TemplateLibrary(template_library_path)
        self.converters: Dict[str, TemplateBasedConverter] = {}
        
        # 预加载所有现有模板
        self._load_existing_templates()
        
        logger.info("高级模板转换器初始化完成")
    
    def _load_existing_templates(self):
        """加载所有现有模板到转换器字典"""
        try:
            templates = self.template_library.list_templates()
            for template_id, template_info in templates.items():
                if 'word_file' in template_info:
                    word_file_path = template_info['word_file']
                    if os.path.exists(word_file_path):
                        try:
                            self.converters[template_id] = TemplateBasedConverter(word_file_path)
                            logger.info(f"已加载模板转换器: {template_id} - {template_info.get('name', 'Unknown')}")
                        except Exception as e:
                            logger.error(f"加载模板转换器失败 {template_id}: {e}")
                    else:
                        logger.warning(f"模板文件不存在: {word_file_path}")
            
            logger.info(f"共加载 {len(self.converters)} 个模板转换器")
        except Exception as e:
            logger.error(f"预加载模板失败: {e}")
    
    def add_template(self, word_file_path: str, template_name: str,
                    description: str = "", tags: List[str] = None) -> str:
        """添加模板到库"""
        template_id = self.template_library.add_template(
            word_file_path, template_name, description, tags
        )
        
        # 创建对应的转换器
        self.converters[template_id] = TemplateBasedConverter(word_file_path)
        
        logger.info(f"模板已添加: {template_id}")
        return template_id
    
    def convert_with_template(self, markdown_content: str, template_id: str,
                             output_path: str) -> ConversionResult:
        """使用指定模板转换"""
        if template_id not in self.converters:
            raise ValueError(f"模板不存在: {template_id}")
        
        converter = self.converters[template_id]
        return converter.convert_markdown_to_word(markdown_content, output_path)
    
    def auto_select_template(self, markdown_content: str) -> Optional[str]:
        """自动选择最适合的模板"""
        document_context = analyze_markdown_for_mapping(markdown_content)
        document_structure = analyze_markdown_document(markdown_content)
        
        # 基于文档类型和特征选择模板
        doc_type = document_context.get('estimated_type', 'general')
        
        # 搜索相关模板
        if doc_type == 'academic':
            templates = self.template_library.search_templates(tags=['academic', 'thesis'])
        elif doc_type == 'technical':
            templates = self.template_library.search_templates(tags=['technical', 'api'])
        elif doc_type == 'business':
            templates = self.template_library.search_templates(tags=['business', 'report'])
        else:
            templates = self.template_library.search_templates(tags=['general'])
        
        # 选择第一个匹配的模板
        if templates:
            return list(templates.keys())[0]
        
        # 如果没有匹配的，返回第一个可用模板
        all_templates = self.template_library.list_templates()
        if all_templates:
            return list(all_templates.keys())[0]
        
        return None
    
    def convert_with_auto_template(self, markdown_content: str,
                                  output_path: str) -> ConversionResult:
        """自动选择模板并转换"""
        template_id = self.auto_select_template(markdown_content)
        
        if not template_id:
            raise ValueError("没有可用的模板")
        
        logger.info(f"自动选择模板: {template_id}")
        return self.convert_with_template(markdown_content, template_id, output_path)
    
    def batch_convert(self, input_files: List[str], output_dir: str,
                     template_id: Optional[str] = None) -> Dict[str, ConversionResult]:
        """批量转换"""
        results = {}
        output_path = Path(output_dir)
        output_path.mkdir(parents=True, exist_ok=True)
        
        for input_file in input_files:
            try:
                # 读取Markdown文件
                with open(input_file, 'r', encoding='utf-8') as f:
                    markdown_content = f.read()
                
                # 生成输出文件名
                input_path = Path(input_file)
                output_file = output_path / f"{input_path.stem}.docx"
                
                # 转换
                if template_id:
                    result = self.convert_with_template(
                        markdown_content, template_id, str(output_file)
                    )
                else:
                    result = self.convert_with_auto_template(
                        markdown_content, str(output_file)
                    )
                
                results[input_file] = result
                
            except Exception as e:
                error_result = ConversionResult(
                    success=False,
                    output_path="",
                    template_used="",
                    style_mappings_applied=0,
                    warnings=[],
                    errors=[str(e)],
                    processing_time=0.0
                )
                results[input_file] = error_result
                logger.error(f"转换文件 {input_file} 失败: {e}")
        
        return results
    
    def get_available_templates(self) -> Dict[str, Any]:
        """获取可用模板列表"""
        return self.template_library.list_templates()
    
    def get_template_info(self, template_id: str) -> Optional[Dict[str, Any]]:
        """获取模板详细信息"""
        if template_id in self.template_library.list_templates():
            template_info = self.template_library.get_template_info(template_id)
            if template_info:
                return template_info.to_dict()
        return None


# 便捷函数
def convert_markdown_with_template(markdown_content: str, word_template_path: str,
                                  output_path: str) -> ConversionResult:
    """使用Word模板转换Markdown的便捷函数"""
    converter = TemplateBasedConverter(word_template_path)
    return converter.convert_markdown_to_word(markdown_content, output_path)


def create_advanced_converter(template_library_path: str = "template_library") -> AdvancedTemplateConverter:
    """创建高级转换器的便捷函数"""
    return AdvancedTemplateConverter(template_library_path)