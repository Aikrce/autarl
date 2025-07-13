#!/usr/bin/env python3
"""
Enhanced Style Engine
优化的样式引擎，提供高性能的样式应用和管理功能
"""

import logging
from typing import Dict, List, Optional, Any, Union
from dataclasses import dataclass
from enum import Enum
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Length
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.styles.style import _ParagraphStyle, _CharacterStyle
import threading
from functools import lru_cache

from enhanced_templates_config import TemplateConfig, StyleConfig, FontConfig, ParagraphConfig

logger = logging.getLogger(__name__)


class AlignmentType(Enum):
    """对齐方式枚举"""
    LEFT = "left"
    CENTER = "center"
    RIGHT = "right"
    JUSTIFY = "justify"


@dataclass
class StyleCache:
    """样式缓存"""
    docx_style: Any
    config: StyleConfig
    last_used: float
    
    def __post_init__(self):
        import time
        if not hasattr(self, 'last_used'):
            self.last_used = time.time()


class EnhancedStyleEngine:
    """增强的样式引擎"""
    
    def __init__(self, template_config: TemplateConfig):
        self.template_config = template_config
        self.document: Optional[Document] = None
        self._style_cache: Dict[str, StyleCache] = {}
        self._lock = threading.RLock()
        self._alignment_map = self._build_alignment_map()
        
        # 性能优化设置
        self.enable_cache = True
        self.max_cache_size = 100
        
        logger.info(f"样式引擎初始化完成，模板: {template_config.name}")
    
    def _build_alignment_map(self) -> Dict[str, Any]:
        """构建对齐方式映射"""
        return {
            'left': WD_PARAGRAPH_ALIGNMENT.LEFT,
            'center': WD_PARAGRAPH_ALIGNMENT.CENTER,
            'right': WD_PARAGRAPH_ALIGNMENT.RIGHT,
            'justify': WD_PARAGRAPH_ALIGNMENT.JUSTIFY,
            'distribute': WD_PARAGRAPH_ALIGNMENT.DISTRIBUTE
        }
    
    def set_document(self, document: Document):
        """设置当前文档"""
        self.document = document
        self._apply_document_settings()
        self._register_template_styles()
    
    def _apply_document_settings(self):
        """应用文档级设置"""
        if not self.document:
            return
        
        # 设置页边距
        margins = self.template_config.page_margins_cm
        for section in self.document.sections:
            section.top_margin = Cm(margins['top'])
            section.bottom_margin = Cm(margins['bottom'])
            section.left_margin = Cm(margins['left'])
            section.right_margin = Cm(margins['right'])
            
            # 设置页面方向
            if self.template_config.page_orientation == 'landscape':
                section.orientation = 1  # WD_ORIENT.LANDSCAPE
            
        logger.info("文档页面设置已应用")
    
    def _register_template_styles(self):
        """注册模板中定义的样式"""
        if not self.document:
            return
        
        styles = self.document.styles
        
        for style_config in self.template_config.styles:
            try:
                # 检查样式是否已存在
                if style_config.name in [s.name for s in styles]:
                    logger.debug(f"样式 '{style_config.name}' 已存在，跳过创建")
                    continue
                
                # 确定样式类型
                style_type = self._get_style_type(style_config.style_type)
                
                # 创建样式
                docx_style = styles.add_style(style_config.name, style_type)
                
                # 应用样式配置
                self._apply_style_config(docx_style, style_config)
                
                # 缓存样式
                if self.enable_cache:
                    self._cache_style(style_config.name, docx_style, style_config)
                
                logger.debug(f"样式 '{style_config.name}' 创建成功")
                
            except Exception as e:
                logger.warning(f"创建样式 '{style_config.name}' 失败: {e}")
    
    def _get_style_type(self, style_type_str: str) -> Any:
        """获取Word样式类型"""
        type_map = {
            'paragraph': WD_STYLE_TYPE.PARAGRAPH,
            'character': WD_STYLE_TYPE.CHARACTER,
            'table': WD_STYLE_TYPE.TABLE,
            'list': WD_STYLE_TYPE.LIST
        }
        return type_map.get(style_type_str, WD_STYLE_TYPE.PARAGRAPH)
    
    def _apply_style_config(self, docx_style: Any, style_config: StyleConfig):
        """应用样式配置到Word样式对象"""
        # 应用字体配置
        if style_config.font and hasattr(docx_style, 'font'):
            self._apply_font_config(docx_style.font, style_config.font)
        
        # 应用段落配置
        if style_config.paragraph and hasattr(docx_style, 'paragraph_format'):
            self._apply_paragraph_config(docx_style.paragraph_format, style_config.paragraph)
        
        # 设置基础样式
        if style_config.based_on:
            try:
                docx_style.base_style = self.document.styles[style_config.based_on]
            except KeyError:
                logger.warning(f"基础样式 '{style_config.based_on}' 不存在")
        
        # 设置下一个样式
        if style_config.next_style:
            try:
                docx_style.next_paragraph_style = self.document.styles[style_config.next_style]
            except KeyError:
                logger.warning(f"下一个样式 '{style_config.next_style}' 不存在")
    
    def _apply_font_config(self, font_obj: Any, font_config: FontConfig):
        """应用字体配置"""
        font_obj.name = font_config.name
        font_obj.size = Pt(font_config.size_pt)
        font_obj.bold = font_config.bold
        font_obj.italic = font_config.italic
        
        # 设置字体颜色
        if font_config.color_rgb != (0, 0, 0):
            font_obj.color.rgb = RGBColor(*font_config.color_rgb)
    
    def _apply_paragraph_config(self, para_format: Any, para_config: ParagraphConfig):
        """应用段落配置"""
        # 设置对齐方式
        if para_config.alignment in self._alignment_map:
            para_format.alignment = self._alignment_map[para_config.alignment]
        
        # 设置间距
        if para_config.space_before_pt > 0:
            para_format.space_before = Pt(para_config.space_before_pt)
        if para_config.space_after_pt > 0:
            para_format.space_after = Pt(para_config.space_after_pt)
        
        # 设置行距
        if para_config.line_spacing != 1.0:
            para_format.line_spacing = para_config.line_spacing
        
        # 设置缩进
        if para_config.first_line_indent_cm > 0:
            para_format.first_line_indent = Cm(para_config.first_line_indent_cm)
        if para_config.left_indent_cm > 0:
            para_format.left_indent = Cm(para_config.left_indent_cm)
        if para_config.right_indent_cm > 0:
            para_format.right_indent = Cm(para_config.right_indent_cm)
        if para_config.hanging_indent_cm > 0:
            para_format.hanging_indent = Cm(para_config.hanging_indent_cm)
    
    def _cache_style(self, name: str, docx_style: Any, config: StyleConfig):
        """缓存样式"""
        with self._lock:
            if len(self._style_cache) >= self.max_cache_size:
                self._cleanup_cache()
            
            self._style_cache[name] = StyleCache(
                docx_style=docx_style,
                config=config,
                last_used=0
            )
    
    def _cleanup_cache(self):
        """清理样式缓存"""
        # 移除最少使用的样式
        if self._style_cache:
            oldest_key = min(self._style_cache.keys(), 
                           key=lambda k: self._style_cache[k].last_used)
            del self._style_cache[oldest_key]
            logger.debug(f"样式缓存清理：移除 '{oldest_key}'")
    
    @lru_cache(maxsize=128)
    def get_style(self, style_name: str) -> Optional[Any]:
        """获取样式对象（带缓存）"""
        if not self.document:
            logger.warning("文档未设置，无法获取样式")
            return None
        
        # 检查缓存
        if self.enable_cache and style_name in self._style_cache:
            import time
            cache_entry = self._style_cache[style_name]
            cache_entry.last_used = time.time()
            return cache_entry.docx_style
        
        # 从文档中获取
        try:
            return self.document.styles[style_name]
        except KeyError:
            logger.warning(f"样式 '{style_name}' 不存在")
            return None
    
    def apply_style_to_paragraph(self, paragraph: Any, style_name: str, 
                                custom_formatting: Optional[Dict[str, Any]] = None) -> bool:
        """将样式应用到段落"""
        try:
            style = self.get_style(style_name)
            if style:
                paragraph.style = style
                
                # 应用自定义格式
                if custom_formatting:
                    self._apply_custom_formatting(paragraph, custom_formatting)
                
                return True
            return False
        except Exception as e:
            logger.error(f"应用段落样式失败: {e}")
            return False
    
    def apply_style_to_run(self, run: Any, style_name: str,
                          custom_formatting: Optional[Dict[str, Any]] = None) -> bool:
        """将样式应用到文本运行"""
        try:
            style = self.get_style(style_name)
            if style and hasattr(style, 'type') and style.type == WD_STYLE_TYPE.CHARACTER:
                run.style = style
                
                # 应用自定义格式
                if custom_formatting:
                    self._apply_custom_run_formatting(run, custom_formatting)
                
                return True
            return False
        except Exception as e:
            logger.error(f"应用文本样式失败: {e}")
            return False
    
    def _apply_custom_formatting(self, paragraph: Any, formatting: Dict[str, Any]):
        """应用自定义段落格式"""
        para_format = paragraph.paragraph_format
        
        # 字体格式（应用到段落中的所有运行）
        if 'font' in formatting:
            font_config = formatting['font']
            for run in paragraph.runs:
                if 'name' in font_config:
                    run.font.name = font_config['name']
                if 'size_pt' in font_config:
                    run.font.size = Pt(font_config['size_pt'])
                if 'bold' in font_config:
                    run.font.bold = font_config['bold']
                if 'italic' in font_config:
                    run.font.italic = font_config['italic']
                if 'color_rgb' in font_config:
                    run.font.color.rgb = RGBColor(*font_config['color_rgb'])
        
        # 段落格式
        if 'alignment' in formatting:
            alignment = formatting['alignment']
            if alignment in self._alignment_map:
                para_format.alignment = self._alignment_map[alignment]
        
        if 'space_before_pt' in formatting:
            para_format.space_before = Pt(formatting['space_before_pt'])
        
        if 'space_after_pt' in formatting:
            para_format.space_after = Pt(formatting['space_after_pt'])
        
        if 'line_spacing' in formatting:
            para_format.line_spacing = formatting['line_spacing']
        
        if 'first_line_indent_cm' in formatting:
            para_format.first_line_indent = Cm(formatting['first_line_indent_cm'])
        
        if 'left_indent_cm' in formatting:
            para_format.left_indent = Cm(formatting['left_indent_cm'])
        
        if 'hanging_indent_cm' in formatting:
            para_format.hanging_indent = Cm(formatting['hanging_indent_cm'])
    
    def _apply_custom_run_formatting(self, run: Any, formatting: Dict[str, Any]):
        """应用自定义文本运行格式"""
        if 'font' in formatting:
            font_config = formatting['font']
            if 'name' in font_config:
                run.font.name = font_config['name']
            if 'size_pt' in font_config:
                run.font.size = Pt(font_config['size_pt'])
            if 'bold' in font_config:
                run.font.bold = font_config['bold']
            if 'italic' in font_config:
                run.font.italic = font_config['italic']
            if 'color_rgb' in font_config:
                run.font.color.rgb = RGBColor(*font_config['color_rgb'])
    
    def create_dynamic_style(self, name: str, base_style: str, 
                           modifications: Dict[str, Any]) -> bool:
        """动态创建样式"""
        if not self.document:
            return False
        
        try:
            # 获取基础样式配置
            base_config = None
            for style_config in self.template_config.styles:
                if style_config.name == base_style:
                    base_config = style_config
                    break
            
            if not base_config:
                logger.warning(f"基础样式 '{base_style}' 不存在")
                return False
            
            # 创建修改后的配置
            import copy
            new_config = copy.deepcopy(base_config)
            new_config.name = name
            
            # 应用修改
            self._apply_modifications_to_config(new_config, modifications)
            
            # 创建Word样式
            style_type = self._get_style_type(new_config.style_type)
            docx_style = self.document.styles.add_style(name, style_type)
            self._apply_style_config(docx_style, new_config)
            
            # 缓存新样式
            if self.enable_cache:
                self._cache_style(name, docx_style, new_config)
            
            logger.info(f"动态样式 '{name}' 创建成功")
            return True
            
        except Exception as e:
            logger.error(f"创建动态样式失败: {e}")
            return False
    
    def _apply_modifications_to_config(self, config: StyleConfig, modifications: Dict[str, Any]):
        """将修改应用到样式配置"""
        if 'font' in modifications:
            font_mods = modifications['font']
            for key, value in font_mods.items():
                if hasattr(config.font, key):
                    setattr(config.font, key, value)
        
        if 'paragraph' in modifications:
            para_mods = modifications['paragraph']
            for key, value in para_mods.items():
                if hasattr(config.paragraph, key):
                    setattr(config.paragraph, key, value)
    
    def add_borders_to_paragraph(self, paragraph: Any, border_config: Dict[str, Any]):
        """为段落添加边框"""
        try:
            pPr = paragraph._p.get_or_add_pPr()
            borders = OxmlElement('w:pBdr')
            
            # 边框位置
            positions = border_config.get('positions', ['top', 'left', 'bottom', 'right'])
            
            for position in positions:
                border = OxmlElement(f'w:{position}')
                border.set(qn('w:val'), border_config.get('style', 'single'))
                border.set(qn('w:sz'), str(border_config.get('width', 4)))
                border.set(qn('w:space'), str(border_config.get('space', 1)))
                border.set(qn('w:color'), border_config.get('color', 'auto'))
                borders.append(border)
            
            pPr.insert_element_before(borders,
                'w:shd', 'w:tabs', 'w:suppressAutoHyphens', 'w:kinsoku',
                'w:wordWrap', 'w:overflowPunct', 'w:topLinePunct', 'w:autoSpaceDE',
                'w:autoSpaceDN', 'w:bidi', 'w:adjustRightInd', 'w:snapToGrid',
                'w:spacing', 'w:ind', 'w:contextualSpacing', 'w:mirrorIndents',
                'w:suppressOverlap', 'w:jc', 'w:textDirection', 'w:textAlignment',
                'w:textboxTightWrap', 'w:outlineLvl', 'w:divId', 'w:cnfStyle',
                'w:rPr', 'w:sectPr', 'w:pPrChange'
            )
            
            return True
        except Exception as e:
            logger.error(f"添加段落边框失败: {e}")
            return False
    
    def setup_header_footer(self):
        """设置页眉页脚"""
        if not self.document:
            return
        
        header_config = self.template_config.header_config
        footer_config = self.template_config.footer_config
        
        for section in self.document.sections:
            # 设置页眉
            if header_config:
                self._setup_header(section, header_config)
            
            # 设置页脚
            if footer_config:
                self._setup_footer(section, footer_config)
    
    def _setup_header(self, section: Any, config: Dict[str, Any]):
        """设置页眉"""
        try:
            header = section.header
            if header.paragraphs:
                header_para = header.paragraphs[0]
                header_para.text = config.get('text', '')
                
                # 设置对齐
                alignment = config.get('alignment', 'center')
                if alignment in self._alignment_map:
                    header_para.alignment = self._alignment_map[alignment]
                
                # 设置字体
                font_config = config.get('font', {})
                for run in header_para.runs:
                    if 'name' in font_config:
                        run.font.name = font_config['name']
                    if 'size_pt' in font_config:
                        run.font.size = Pt(font_config['size_pt'])
                    if 'bold' in font_config:
                        run.font.bold = font_config['bold']
        except Exception as e:
            logger.error(f"设置页眉失败: {e}")
    
    def _setup_footer(self, section: Any, config: Dict[str, Any]):
        """设置页脚"""
        try:
            footer = section.footer
            if footer.paragraphs and config.get('show_page_number', False):
                footer_para = footer.paragraphs[0]
                footer_para.clear()
                
                # 设置对齐
                alignment = config.get('alignment', 'center')
                if alignment in self._alignment_map:
                    footer_para.alignment = self._alignment_map[alignment]
                
                # 添加页码
                run = footer_para.add_run()
                font_config = config.get('font', {})
                if 'name' in font_config:
                    run.font.name = font_config['name']
                if 'size_pt' in font_config:
                    run.font.size = Pt(font_config['size_pt'])
                
                # 插入页码字段
                self._insert_page_number_field(run)
        except Exception as e:
            logger.error(f"设置页脚失败: {e}")
    
    def _insert_page_number_field(self, run: Any):
        """插入页码字段"""
        try:
            fldChar1 = OxmlElement('w:fldChar')
            fldChar1.set(qn('w:fldCharType'), 'begin')
            run._element.append(fldChar1)
            
            instrText = OxmlElement('w:instrText')
            instrText.text = 'PAGE'
            run._element.append(instrText)
            
            fldChar2 = OxmlElement('w:fldChar')
            fldChar2.set(qn('w:fldCharType'), 'end')
            run._element.append(fldChar2)
        except Exception as e:
            logger.error(f"插入页码字段失败: {e}")
    
    def get_style_statistics(self) -> Dict[str, Any]:
        """获取样式统计信息"""
        stats = {
            'total_styles': len(self.template_config.styles),
            'cached_styles': len(self._style_cache),
            'style_types': {},
            'memory_usage': 0
        }
        
        # 统计样式类型
        for style_config in self.template_config.styles:
            style_type = style_config.style_type
            if style_type not in stats['style_types']:
                stats['style_types'][style_type] = 0
            stats['style_types'][style_type] += 1
        
        # 估算内存使用
        import sys
        for cache_entry in self._style_cache.values():
            stats['memory_usage'] += sys.getsizeof(cache_entry)
        
        return stats
    
    def optimize_performance(self):
        """优化性能设置"""
        # 清理缓存
        self._cleanup_cache()
        
        # 清理LRU缓存
        self.get_style.cache_clear()
        
        logger.info("样式引擎性能优化完成")
    
    def export_styles_to_dict(self) -> Dict[str, Any]:
        """导出样式配置为字典"""
        return {
            'template_name': self.template_config.name,
            'styles': [style.to_dict() for style in self.template_config.styles],
            'statistics': self.get_style_statistics()
        }
    
    def batch_apply_styles(self, style_applications: List[Dict[str, Any]]) -> Dict[str, bool]:
        """批量应用样式"""
        results = {}
        
        for application in style_applications:
            element = application.get('element')
            style_name = application.get('style_name')
            custom_formatting = application.get('custom_formatting')
            element_type = application.get('type', 'paragraph')
            
            if element_type == 'paragraph':
                success = self.apply_style_to_paragraph(element, style_name, custom_formatting)
            elif element_type == 'run':
                success = self.apply_style_to_run(element, style_name, custom_formatting)
            else:
                success = False
            
            results[f"{element_type}_{id(element)}"] = success
        
        logger.info(f"批量样式应用完成: {sum(results.values())}/{len(results)} 成功")
        return results


class StyleEngineFactory:
    """样式引擎工厂"""
    
    @staticmethod
    def create_engine(template_config: TemplateConfig) -> EnhancedStyleEngine:
        """创建样式引擎实例"""
        return EnhancedStyleEngine(template_config)
    
    @staticmethod
    def create_optimized_engine(template_config: TemplateConfig, 
                              cache_size: int = 200,
                              enable_performance_mode: bool = True) -> EnhancedStyleEngine:
        """创建优化的样式引擎实例"""
        engine = EnhancedStyleEngine(template_config)
        engine.max_cache_size = cache_size
        
        if enable_performance_mode:
            # 预加载常用样式
            engine._preload_common_styles()
        
        return engine
    
    def _preload_common_styles(self):
        """预加载常用样式"""
        # 这里可以预加载一些常用样式到缓存中
        pass